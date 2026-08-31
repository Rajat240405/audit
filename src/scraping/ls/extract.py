"""Document bytes → (question_text, answer_text) extraction — LS primary text stage.

Lok Sabha records have NO inline text (the frozen workbook's
``questionText``/``answerText``/``answerTextHindi`` columns are entirely
empty), so extraction from the official answer documents is the primary
content source — not a fallback.

Reuses the canonical, production-proven stack from the legacy archive scraper
(``src/data/scraper.py``, untouched):

- ``RealArchiveScraper._extract_pdf_text_bytes`` — PyMuPDF table-aware first,
  pypdf fallback.
- ``RealArchiveScraper._split_question_answer`` — ANSWER/REPLY boundary
  split (with the legacy ratio fallback).

The DOCX paragraph/table walk is re-implemented here (``_docx_text``) because
the legacy copy is an instance method — the equivalence is pinned by a test
against the legacy method. Nothing is ever synthesized: failures return a
machine-readable reason and the record keeps empty text.
"""

from __future__ import annotations

import io

from src.data.scraper import RealArchiveScraper

#: extraction failure reasons (verbatim legacy vocabulary)
Reason = str  # "scanned" | "parser_failure" | "unsupported"


def _docx_text(data: bytes) -> str | None:
    """All text of an OOXML DOCX (paragraphs + table cells), mirroring
    ``RealArchiveScraper._extract_text_from_docx`` (equivalence-tested)."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover — environment guard
        raise RuntimeError(
            "python-docx is required for DOCX support. Install with: pip install python-docx"
        ) from exc
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:  # noqa: BLE001
        return None


def extract_qa(body: bytes, doc_format: str) -> tuple[tuple[str, str] | None, Reason | None]:
    """Extract (question, answer) from document bytes of a sniffed format.

    Returns ``((question, answer), None)`` on success, else ``(None, reason)``
    with reason in the legacy vocabulary. Empty body/format values degrade to
    ``unsupported`` — callers never see an exception from this stage.
    """
    if not body:
        return None, "unsupported"
    if doc_format == "pdf":
        text = RealArchiveScraper._extract_pdf_text_bytes(body)  # noqa: SLF001 — canonical reuse
        if text is None:
            return None, "parser_failure"
        if not text.strip():
            return None, "scanned"
    elif doc_format == "docx":
        text = _docx_text(body)
        if text is None:
            return None, "parser_failure"
        if not text.strip():
            return None, "scanned"
    else:
        return None, "unsupported"
    return RealArchiveScraper._split_question_answer(text), None  # noqa: SLF001
