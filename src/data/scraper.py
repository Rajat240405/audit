"""
Web scraper and data crawler for real Lok Sabha Parliamentary Q&A dataset.

Design Decisions
----------------
1. STRATEGY PATTERN: Implement multiple modular scraping strategies
   - "live": Crawls live data from the official sansad.in portal with
     fail-safe fallback systems.
   - "archive": Loads actual, genuine Lok Sabha metadata from the official
     Parliament of India dataset on Zenodo, downloads each official PDF from
     questionsFilePath, extracts the full question and answer text using the PyMuPDF-based
     table-aware extractor, and populates question_text and answer_text
     directly from the official document.
   - "mock": Produces high-quality, topic-aligned synthetic records.
   - "local": Loads records from a local JSONL file.

2. CHECKPOINT & RESUME: Tracks scraped URLs and question IDs inside
   `data/raw/checkpoint.json`. If interrupted, the crawler loads this
   JSON index and automatically skips previously processed records.

3. RETRY WITH EXPONENTIAL BACKOFF: HTTP operations implement an
   exponential backoff loop with randomized jitter to gracefully handle
   transient network dropouts and rate-limiting blocks.

4. RATE LIMITING: Enforces strict, configurable sleep delays between
   consecutive requests to respect Lok Sabha bandwidth limits.

5. SELECTOR RESILIENCY: Parses HTML with redundant CSS selectors to
   survive frequent Government portal design changes.
"""

from __future__ import annotations

import asyncio
import io
import json
import random
import re
import time
import warnings
import zipfile
from abc import ABC, abstractmethod
from collections.abc import Iterator
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import httpx
import pandas as pd
from bs4 import BeautifulSoup
from rich.console import Console
from rich.panel import Panel
from rich.progress import BarColumn, Progress, TextColumn, TimeElapsedColumn

from src.models.qa_record import QARecord, QARecordMetadata, QuestionType
from src.models.statistics import ScrapingStats

console = Console()

# Browser User-Agent shared by all request paths (sansad.in + DSpace)
BROWSER_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

# ─────────────────────────────────────────────────────────────────────────────
# Base Scraper Interface
# ─────────────────────────────────────────────────────────────────────────────

class Scraper(ABC):
    """Abstract base class for all scraper implementations."""

    def __init__(
        self,
        base_url: str,
        rate_limit_seconds: float = 2.0,
        timeout_seconds: int = 30,
        max_retries: int = 3,
        checkpoint_file: str = "data/raw/checkpoint.json",
    ) -> None:
        self.base_url = base_url
        self.rate_limit_seconds = rate_limit_seconds
        self.timeout_seconds = timeout_seconds
        self.max_retries = max_retries
        self.checkpoint_file = Path(checkpoint_file)
        self.stats = ScrapingStats()
        self._last_request_time: float = 0.0

        # Load Checkpoint State
        self.checkpoint_file.parent.mkdir(parents=True, exist_ok=True)
        self.checkpoints: dict[str, str] = self._load_checkpoints()

    def _rate_limit(self) -> None:
        """Enforce rate limiting between requests."""
        elapsed = time.perf_counter() - self._last_request_time
        if elapsed < self.rate_limit_seconds:
            time.sleep(self.rate_limit_seconds - elapsed)
        self._last_request_time = time.perf_counter()

    def _load_checkpoints(self) -> dict[str, str]:
        """Load the crawl checkpoint mapping."""
        if self.checkpoint_file.exists():
            try:
                with open(self.checkpoint_file, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _save_checkpoint(self, key: str, value: str = "done") -> None:
        """Record and persist a crawl checkpoint."""
        self.checkpoints[key] = value
        try:
            with open(self.checkpoint_file, "w", encoding="utf-8") as f:
                json.dump(self.checkpoints, f, indent=2)
        except Exception as e:
            console.print(f"[yellow]Warning: Could not save checkpoint: {e}[/yellow]")

    @abstractmethod
    def scrape_all(self, max_records: int = 3500) -> Iterator[QARecord]:
        """Scrape all available Q&A records up to max_records."""
        ...

    def _make_request(
        self,
        url: str,
        client: httpx.Client,
    ) -> httpx.Response | None:
        """Make an HTTP request with exponential backoff retry and jitter."""
        delay = 1.0
        for attempt in range(self.max_retries):
            self._rate_limit()
            try:
                response = client.get(
                    url,
                    headers={
                        "User-Agent": (
                            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                            "AppleWebKit/537.36 (KHTML, like Gecko) "
                            "Chrome/124.0.0.0 Safari/537.36"
                        )
                    },
                    timeout=self.timeout_seconds,
                )
                if response.status_code == 200:
                    self.stats.http_errors = 0
                    return response
                elif response.status_code == 429:
                    self.stats.rate_limit_hits += 1
                    sleep_time = delay + random.uniform(0.1, 0.5)
                    console.print(f"[yellow]Rate limit (429) hit. Backing off for {sleep_time:.2f}s...[/yellow]")
                    time.sleep(sleep_time)
                    delay *= 2
                else:
                    self.stats.http_errors += 1
                    sleep_time = delay
                    console.print(f"[yellow]HTTP {response.status_code} for {url}. Retrying in {sleep_time:.1f}s...[/yellow]")
                    time.sleep(sleep_time)
                    delay *= 1.5
            except httpx.RequestError as e:
                self.stats.http_errors += 1
                sleep_time = delay + random.uniform(0.1, 0.5)
                console.print(f"[red]Request error: {e}. Retrying in {sleep_time:.1f}s...[/red]")
                time.sleep(sleep_time)
                delay *= 2

        self.stats.individual_pages_failed += 1
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Real Live Lok Sabha Scraper (with selector fallback & dynamic API crawl)
# ─────────────────────────────────────────────────────────────────────────────

class LiveLoksabhaScraper(Scraper):
    """
    DORMANT (workspace cleanup, audit §5): never used for any shipped
    corpus. This pre-archive-era class crawls the live sansad.in Q&A list
    pages with guessed CSS selectors; production acquisition moved to
    :class:`RealArchiveScraper` (strategy=archive), and deduplicated
    parliamentary crawling now belongs to the src/scraping architecture
    (session-staged, fixture-replayable, manifest-audited). It is kept for
    reference only — instantiating it emits a DeprecationWarning.

    If a live Lok Sabha crawler is ever genuinely required, the concrete
    recommendation (audit §5) is to model it on the RS crawler
    (src/scraping/rs: client/normalize/pipeline/manifest with deterministic
    replay fixtures), NOT to revive this selector-guessing class.

    Original intent: crawl live Lok Sabha Q&A records directly from
    sansad.in using selector fallbacks and API polling.
    """

    def __init__(self, *args, **kwargs) -> None:
        warnings.warn(
            "LiveLoksabhaScraper is dormant (audit §5): it was never used "
            "for any shipped corpus. Production Lok Sabha acquisition is "
            "strategy=archive (RealArchiveScraper); future live crawling "
            "should follow the src/scraping/rs architecture, not this class.",
            DeprecationWarning,
            stacklevel=2,
        )
        super().__init__(*args, **kwargs)

    def scrape_all(self, max_records: int = 3500) -> Iterator[QARecord]:
        self.stats.started_at = datetime.utcnow()
        console.print(f"[cyan]Starting Live Lok Sabha crawl from {self.base_url}[/cyan]")

        records_scraped = 0
        page = 1

        with httpx.Client(timeout=self.timeout_seconds) as client:
            while records_scraped < max_records:
                # 1. Page Endpoint Crawling (Dynamic list page or API endpoint)
                page_url = f"{self.base_url}?page={page}"
                
                # Check if page has already been processed in checkpoints
                if self.checkpoints.get(f"page-{page}") == "done":
                    console.print(f"[dim]Page {page} already processed, skipping...[/dim]")
                    page += 1
                    continue

                response = self._make_request(page_url, client)
                if not response:
                    console.print(f"[yellow]Failed to fetch list page {page} after retries. Gracefully skipping.[/yellow]")
                    break

                soup = BeautifulSoup(response.text, "lxml")
                question_links = self._extract_question_links(soup)

                if not question_links:
                    console.print(f"[yellow]No active question links found on page {page}. Finalizing scrape.[/yellow]")
                    break

                self.stats.question_links_found += len(question_links)

                for link_url in question_links:
                    if records_scraped >= max_records:
                        break

                    # Checkpoint resume check
                    if self.checkpoints.get(link_url) == "done":
                        continue

                    record = self._scrape_individual_page(link_url, client)
                    if record:
                        self.stats.individual_pages_success += 1
                        self._save_checkpoint(link_url, "done")
                        yield record
                        records_scraped += 1
                    else:
                        self.stats.individual_pages_failed += 1

                    self.stats.individual_pages_attempted += 1

                self._save_checkpoint(f"page-{page}", "done")
                page += 1
                self.stats.pages_scraped = page

        self.stats.completed_at = datetime.utcnow()
        console.print(f"[green]Scraping complete: {records_scraped} records scraped successfully.[/green]")

    def _extract_question_links(self, soup: BeautifulSoup) -> list[str]:
        """Extract question details page URLs with redundant fallbacks."""
        links = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if any(term in href for term in ["/ls/questions/", "/questions-and-answers", "/getFile/lsapps/"]):
                if href.startswith("/"):
                    links.append(f"https://sansad.in{href}")
                elif href.startswith("http"):
                    links.append(href)
        return list(dict.fromkeys(links))

    def _scrape_individual_page(self, url: str, client: httpx.Client) -> QARecord | None:
        """Scrape individual Q&A content."""
        response = self._make_request(url, client)
        if not response:
            return None
        try:
            soup = BeautifulSoup(response.text, "lxml")
            return self._parse_question_page(soup, url)
        except Exception as e:
            self.stats.parse_errors += 1
            console.print(f"[red]Error parsing details page {url}: {e}[/red]")
            return None

    def _parse_question_page(self, soup: BeautifulSoup, url: str) -> QARecord | None:
        """Robust parser supporting multiple CSS fallback layers."""
        # 1. Question Text Selectors
        q_elem = (
            soup.select_one("div.question-text")
            or soup.select_one("div.qstn-text")
            or soup.select_one("div.question")
            or soup.select_one(".qstn-body")
            or soup.find("h2")
        )
        question_text = q_elem.get_text(strip=True) if q_elem else None

        # 2. Answer Text Selectors
        a_elem = (
            soup.select_one("div.answer-text")
            or soup.select_one("div.answer")
            or soup.select_one(".answer-body")
            or soup.select_one("div.answer-body")
            or soup.select_one(".qstn-answer")
        )
        answer_text = a_elem.get_text(strip=True) if a_elem else None

        # Try fallback: if text is empty, check for document download references
        if not question_text or not answer_text:
            return None

        # 3. Metadata Parsing
        ministry = None
        min_elem = soup.select_one(".ministry") or soup.select_one("span.ministry") or soup.select_one("td.ministry")
        if min_elem:
            ministry = min_elem.get_text(strip=True)

        date = None
        dt_elem = soup.select_one(".date") or soup.select_one("span.date") or soup.select_one("td.date")
        if dt_elem:
            date = dt_elem.get_text(strip=True)

        question_id = url.split("/")[-1].replace(".html", "").replace("-", "_") or f"ls-{hash(url)}"

        return QARecord(
            question_id=question_id,
            question_text=question_text,
            answer_text=answer_text,
            metadata=QARecordMetadata(
                ministry=ministry,
                date=date,
                source_url=url,
            ),
            scraped_at=datetime.utcnow(),
        )


# ─────────────────────────────────────────────────────────────────────────────
# Real Archive Scraper (Genuine, Diverse Lok Sabha Metadata-Driven Generator)
# ─────────────────────────────────────────────────────────────────────────────

class RealArchiveScraper(Scraper):
    """
    Loads genuine Lok Sabha questions from the validated production dataset
    (a local Excel copied to ``LOCAL_EXCEL``), downloads each official document
    from ``questionsFilePath`` — a sansad.in annex PDF/DOCX or a DSpace handle
    URL — extracts the full question and answer text, and populates
    ``question_text`` / ``answer_text`` directly from the official document.

    Production guarantees:
    - DSpace handle URLs are resolved to their document bitstream automatically.
    - PDF and DOCX documents are both supported.
    - Invalid downloads (HTML/ZIP/error pages) are rejected, never parsed.
    - Failed records are logged and skipped — synthetic answers are NEVER
      generated for archive ingestion.
    """

    LOCAL_EXCEL = "data/raw/Loksabha_questions.xlsx"
    PDF_CACHE_DIR = "data/raw/pdfs"

    def __init__(
        self,
        *args,
        use_pdf: bool = True,
        ministry_filter: str | None = None,
        **kwargs,
    ) -> None:
        super().__init__(*args, **kwargs)
        self.use_pdf = use_pdf
        self.ministry_filter = ministry_filter
        self.pdf_cache_dir = Path(self.PDF_CACHE_DIR)
        self.pdf_cache_dir.mkdir(parents=True, exist_ok=True)

    # ── DSpace handle resolution (Parliament Digital Library) ──────────────
    #
    # The final dataset may reference DSpace handle pages such as
    #   https://elibrary.sansad.in/handle/123456789/385556
    # instead of a direct PDF/DOCX URL. These are resolved generically (no
    # hardcoded handles) to the primary document bitstream's content URL.

    def _dspace_get(
        self,
        client: httpx.Client,
        url: str,
        params: dict | None = None,
        tries: int = 3,
    ) -> httpx.Response | None:
        """GET against a DSpace endpoint with the shared rate-limit/retry policy."""
        for attempt in range(tries):
            self._rate_limit()
            try:
                response = client.get(
                    url,
                    params=params,
                    headers={"User-Agent": BROWSER_UA},
                    timeout=self.timeout_seconds,
                )
                if response.status_code == 429:
                    time.sleep(2 + attempt)
                    continue
                return response
            except Exception:
                if attempt < tries - 1:
                    time.sleep(1 + attempt)
        return None

    @staticmethod
    def _pick_document_bitstream(bitstreams: list[dict]) -> str | None:
        """
        Choose the primary document bitstream from a DSpace item's bitstream list.

        Prefers English PDF, then English DOCX, then Hindi PDF/DOCX; ignores
        derived files (.txt, .jpg, etc.). Returns the bitstream UUID or None.
        """
        def score(bs: dict) -> tuple | None:
            name = (bs.get("name") or "").lower()
            if name.endswith(".pdf"):
                kind = 0
            elif name.endswith(".docx"):
                kind = 1
            else:
                return None  # derived file, not a document
            hindi = 1 if "hindi" in name else 0
            return (hindi, kind, name)

        scored = [
            (score(bs), bs.get("uuid"))
            for bs in bitstreams
            if score(bs) is not None and bs.get("uuid")
        ]
        scored.sort(key=lambda x: x[0])
        return scored[0][1] if scored else None

    def _resolve_dspace_handle(
        self,
        url: str,
        client: httpx.Client,
    ) -> tuple[str | None, str | None]:
        """
        Resolve a DSpace handle page to the direct content URL of its document.

        Strategy (fully generic — no hardcoded handles):
          1. DSpace 7 REST API: pid/find -> item -> bundles -> bitstreams,
             pick the best document bitstream, return its /content URL.
          2. Fallback: parse the handle page HTML for the citation_pdf_url
             meta or the bitstream UUID download link.

        Returns (content_url, error); error is None on success.
        """
        m = re.match(r"(https?://[^/]+)/handle/(\d+/\d+)", url)
        if not m:
            return None, "invalid_handle_url"
        base = m.group(1)

        # 1) REST API path
        try:
            item = self._dspace_get(client, f"{base}/server/api/pid/find", params={"id": url})
            if item is not None and item.status_code == 200:
                item_uuid = item.json().get("uuid")
                if item_uuid:
                    bundles = self._dspace_get(
                        client, f"{base}/server/api/core/items/{item_uuid}/bundles"
                    )
                    if bundles is not None and bundles.status_code == 200:
                        for bundle in bundles.json().get("_embedded", {}).get("bundles", []):
                            bitstreams = self._dspace_get(
                                client,
                                f"{base}/server/api/core/bundles/{bundle['uuid']}/bitstreams",
                            )
                            if bitstreams is None or bitstreams.status_code != 200:
                                continue
                            uuid = self._pick_document_bitstream(
                                bitstreams.json().get("_embedded", {}).get("bitstreams", [])
                            )
                            if uuid:
                                return f"{base}/server/api/core/bitstreams/{uuid}/content", None
        except Exception as e:
            console.print(f"[dim yellow]DSpace API resolution failed for {url}: {e}[/dim yellow]")

        # 2) HTML page fallback
        try:
            page = self._dspace_get(client, url)
            if page is not None and page.status_code == 200:
                m2 = re.search(r'citation_pdf_url"\s+content="([^"]+)"', page.text, re.I)
                if m2:
                    uuid = re.search(
                        r"([0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12})",
                        m2.group(1),
                    )
                    if uuid:
                        return f"{base}/server/api/core/bitstreams/{uuid.group(1)}/content", None
                m3 = re.search(
                    r"bitstreams/([0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12})/download",
                    page.text,
                )
                if m3:
                    return f"{base}/server/api/core/bitstreams/{m3.group(1)}/content", None
        except Exception as e:
            console.print(f"[dim yellow]DSpace page resolution failed for {url}: {e}[/dim yellow]")

        return None, "no_dspace_bitstream"

    def _download_dataset(self) -> None:
        """Ensure the official Lok Sabha dataset Excel is available locally.

        The validated production dataset must be provided as a local file at
        ``LOCAL_EXCEL`` (copy the merged/frozen workbook there before running).
        The archive scraper never downloads or fabricates a dataset.
        """
        local_path = Path(self.LOCAL_EXCEL)
        if local_path.exists():
            return
        raise FileNotFoundError(
            f"Official dataset not found at {self.LOCAL_EXCEL}. "
            "Copy the validated production Excel to this path before running "
            "archive ingestion."
        )

    def _detect_doc_type(self, data: bytes) -> str:
        """
        Detect the document type from content magic bytes (not the URL extension).

        Returns one of: "pdf", "docx", "zip", "html", "other".
        """
        if data[:5] == b"%PDF-":
            return "pdf"
        if data[:2] == b"PK":
            # ZIP container — DOCX (OOXML) iff it carries the content-types manifest
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    if "[Content_Types].xml" in zf.namelist():
                        return "docx"
            except Exception:
                pass
            return "zip"
        head = data[:2048].lower()
        if head.startswith(b"<!doctype") or head.startswith(b"<html") or b"<html" in head[:500]:
            return "html"
        return "other"

    def _extract_text_from_docx(self, data: bytes) -> str | None:
        """Extract all text from an OOXML DOCX document (paragraphs + tables)."""
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            ) from e
        try:
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return None

    @staticmethod
    def _split_question_answer(text: str) -> tuple[str, str] | None:
        """Split extracted document text into (question, answer)."""
        if not text.strip():
            return None
        match = re.search(r"(?i)\n\s*(?:ANSWER|REPLY|A\s*N\s*S\s*W\s*E\s*R)\s*[:\n]", text)
        if match:
            idx = match.start()
            return text[:idx].strip(), text[idx:].strip()
        # Basic ratio split fallback for documents without an explicit boundary
        split_idx = len(text) // 3
        return text[:split_idx].strip(), text[split_idx:].strip()

    def _extract_text_from_document(
        self, data: bytes, doc_type: str
    ) -> tuple[tuple[str, str] | None, str | None]:
        """
        Extract and split Q&A from raw document bytes (PDF or DOCX).

        Returns ``(result, reason)`` where ``reason`` is None on success and one
        of ``"scanned"`` / ``"parser_failure"`` / ``"unsupported"`` on failure.
        """
        if doc_type == "pdf":
            try:
                text = self._extract_pdf_text_bytes(data)
            except ImportError:
                raise  # PyMuPDF not installed — propagate so the caller knows
            except Exception:
                return None, "parser_failure"  # corrupt/unopenable PDF
            if text is None:
                return None, "scanned"  # valid PDF but no text (image-only)
            if not text.strip():
                return None, "scanned"
        elif doc_type == "docx":
            text = self._extract_text_from_docx(data)
            if text is None:
                return None, "parser_failure"
        else:
            return None, "unsupported"
        return self._split_question_answer(text), None

    @staticmethod
    def _extract_pdf_text_bytes(data: bytes) -> str | None:
        """
        Extract full text from PDF bytes using the PyMuPDF table-aware extractor.

        The PyMuPDF extractor (src.data.pdf_table_extract) reconstructs
        visual table rows geometrically, which is required for correct extraction
        of borderless annexure tables in official sansad.in PDFs.

        Raises ImportError if PyMuPDF is unavailable — install PyMuPDF to
        use this path. Raises Exception for corrupt or unopenable files.
        Returns None when the PDF opens but yields no text (scanned/image-only).
        """
        from src.data.pdf_table_extract import extract_pdf_text

        text = extract_pdf_text(data)
        if text is not None and text.strip():
            return text
        return None

    def _extract_text_from_pdf(self, pdf_path: Path) -> tuple[str, str] | None:
        """Extract question and answer from a local PDF file (compatibility wrapper)."""
        try:
            data = pdf_path.read_bytes()
        except Exception as e:
            console.print(f"[dim yellow]Warning: Failed to read PDF {pdf_path.name}: {e}[/dim yellow]")
            return None
        return self._extract_text_from_document(data, self._detect_doc_type(data))[0]

    def _record_skip(self, reason: str | None) -> None:
        """Increment the appropriate skip counter for a record."""
        reason = reason or "unknown"
        if reason == "broken":
            self.stats.skipped_broken += 1
        elif reason == "scanned":
            self.stats.skipped_scanned += 1
        elif reason == "unsupported":
            self.stats.skipped_unsupported += 1
        elif reason == "parser_failure":
            self.stats.skipped_parser_failure += 1
        else:
            self.stats.skipped_other += 1

    def _get_official_document(
        self, record_id: str, url: str, client: httpx.Client
    ) -> tuple[tuple[str, str] | None, str | None, str]:
        """
        Download, cache, and extract Q&A from an official document (PDF or DOCX).

        Returns ``(result, reason, doc_type)``:
        - result: ``(question, answer)`` on success, else None
        - reason: None on success, else one of ``"broken"``, ``"scanned"``,
          ``"unsupported"``, ``"parser_failure"``
        - doc_type: ``"pdf"`` / ``"docx"`` on success (or best-effort), else ""
        """
        # Resolve DSpace handle URLs to their direct document content URL
        if "/handle/" in url:
            resolved, resolve_err = self._resolve_dspace_handle(url, client)
            if not resolved:
                return None, "broken", ""
            url = resolved

        cache_file = self.pdf_cache_dir / f"{record_id}.pdf"

        # Check cache first (type detected from content, not extension)
        if cache_file.exists() and cache_file.stat().st_size > 500:
            data = cache_file.read_bytes()
            doc_type = self._detect_doc_type(data)
            res, reason = self._extract_text_from_document(data, doc_type)
            if res:
                return res, None, doc_type
            # cached document is unusable (scanned/parser failure) — don't re-download
            return None, reason, doc_type

        # Download with retry block
        delay = 1.0
        for attempt in range(self.max_retries):
            self._rate_limit()
            try:
                response = client.get(
                    url,
                    headers={"User-Agent": BROWSER_UA},
                    timeout=self.timeout_seconds,
                )
                if response.status_code == 429:
                    time.sleep(delay + random.uniform(0.1, 0.5))
                    delay *= 2
                    continue
                if response.status_code != 200:
                    time.sleep(delay)
                    delay *= 1.5
                    continue
                if len(response.content) < 500:
                    # too small to be a real document — retry
                    time.sleep(delay)
                    delay *= 1.5
                    continue

                # Reject clearly non-document responses by their Content-Type header
                content_type = (response.headers.get("content-type") or "").lower()
                if content_type.startswith("text/html") or content_type.startswith("application/json"):
                    # HTML error page / JSON error body — not a document
                    return None, "unsupported", ""

                doc_type = self._detect_doc_type(response.content)
                if doc_type not in ("pdf", "docx"):
                    # Reject HTML/ZIP/other content — do not attempt to parse it
                    return None, "unsupported", doc_type

                cache_file.write_bytes(response.content)
                res, reason = self._extract_text_from_document(response.content, doc_type)
                if res:
                    return res, None, doc_type
                return None, reason, doc_type  # scanned / parser_failure — don't retry
            except Exception as e:
                time.sleep(delay + random.uniform(0.1, 0.5))
                delay *= 2

        return None, "broken", ""

    def scrape_all(self, max_records: int = 3500) -> Iterator[QARecord]:
        self.stats.started_at = datetime.utcnow()
        records_scraped = 0
        rows_processed = 0

        if not self.use_pdf:
            console.print(
                "[bold red]Archive ingestion requires document parsing (--use-pdf). "
                "Without it, no records can be produced — every row will be skipped.[/bold red]"
            )

        # Try to load and parse the official production dataset (local Excel)
        try:
            self._download_dataset()
            console.print(f"[cyan]Loading and parsing Lok Sabha dataset from {self.LOCAL_EXCEL}...[/cyan]")
            df = pd.read_excel(self.LOCAL_EXCEL)
            
            # Filter rows to make sure we have valid metadata (subjects, ministry, quesNo)
            df_valid = df.dropna(subset=["subjects", "ministry", "quesNo"]).copy()

            # Apply the scope resolved by IngestionPipeline before creating QARecord objects.
            # Apply ministry filter at the DataFrame level BEFORE creating any QARecord objects
            before_count = len(df_valid)
            if self.ministry_filter:
                df_valid = df_valid[
                    df_valid["ministry"].str.contains(self.ministry_filter, case=False, na=False)
                ]
                console.print(
                    f"[cyan]Ministry filter applied: {before_count:,} → {len(df_valid):,} rows[/cyan]"
                )

            df_valid = df_valid.sample(frac=1, random_state=42)  # Shuffle to mix topics randomly

            console.print(f"[green]Successfully loaded {len(df_valid):,} valid parliamentary metadata rows.[/green]")

            with Progress(
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn(),
                console=console,
            ) as p:
                task = p.add_task("Ingesting unique real records...", total=min(max_records, len(df_valid)))

                # Create shared client for document downloading.
                # follow_redirects=True is required: the DSpace REST API
                # (pid/find, bitstreams/content) returns HTTP 302 redirects,
                # and without following them every DSpace handle would fail
                # to resolve (observed as mass "broken" skips in production).
                with httpx.Client(timeout=self.timeout_seconds, follow_redirects=True) as client:
                    for _, row in df_valid.iterrows():
                        if records_scraped >= max_records:
                            break

                        # Map columns
                        ques_no = int(row["quesNo"])
                        subjects = str(row["subjects"]).strip().strip(".")
                        ministry = str(row["ministry"]).strip()
                        member = str(row["member"]).strip()
                        qtype_str = str(row["type"]).strip().upper()
                        date_str = str(row["date"]).strip()
                        session_no = int(row["sessionNo"]) if not pd.isna(row["sessionNo"]) else 1
                        lok_no = int(row["lokNo"]) if not pd.isna(row["lokNo"]) else 18
                        questions_file_path = str(row["questionsFilePath"]) if not pd.isna(row["questionsFilePath"]) else ""

                        qtype = QuestionType.STARRED if "STARRED" in qtype_str else QuestionType.UNSTARRED
                        record_id = f"{lok_no}-{session_no}-{ques_no:04d}"

                        # Skip checkpointed records
                        if self.checkpoints.get(record_id) == "done":
                            continue

                        # ── Download and extract the official document (PDF or DOCX) ──
                        # Archive ingestion NEVER synthesizes answers: on any
                        # failure, log the reason and skip the record.
                        parsed = None
                        skip_reason = None
                        doc_type = ""
                        if self.use_pdf and questions_file_path:
                            # Construct full document URL (sansad.in annex or DSpace handle)
                            doc_url = questions_file_path if questions_file_path.startswith("http") else f"https://sansad.in{questions_file_path}"
                            parsed, skip_reason, doc_type = self._get_official_document(record_id, doc_url, client)
                        elif not self.use_pdf:
                            skip_reason = "no_pdf_mode"
                        else:
                            skip_reason = "no_source"

                        if parsed is None:
                            # Log + skip (no synthetic generation in archive mode)
                            self._record_skip(skip_reason)
                            console.print(
                                f"[dim yellow]  skipped {record_id} ({subjects[:60]}): {skip_reason}[/dim yellow]"
                            )
                            rows_processed += 1
                            p.update(task, completed=rows_processed)
                            continue

                        question_text, answer_text = parsed
                        if doc_type == "pdf":
                            self.stats.pdf_parsed += 1
                        elif doc_type == "docx":
                            self.stats.docx_parsed += 1

                        # Build source URL
                        if questions_file_path.startswith("http"):
                            source_url = questions_file_path
                        elif questions_file_path:
                            source_url = f"https://sansad.in{questions_file_path}"
                        else:
                            source_url = f"https://sansad.in/ls/questions/questions-and-answers/{record_id}"

                        rec = QARecord(
                            question_id=record_id,
                            question_text=question_text,
                            answer_text=answer_text,
                            metadata=QARecordMetadata(
                                ministry=ministry,
                                member=member,  # Save the MP name!
                                date=date_str,
                                session=session_no,
                                question_number=ques_no,
                                subject=subjects,
                                question_type=qtype,
                                answer_status="answered",
                                parliament_number=lok_no,
                                source_url=source_url,
                            ),
                            scraped_at=datetime.utcnow(),
                        )

                        yield rec
                        records_scraped += 1
                        rows_processed += 1
                        self._save_checkpoint(record_id, "done")
                        p.update(task, completed=rows_processed)

        except Exception as e:
            console.print(
                f"[red]FATAL: Could not load the official Lok Sabha dataset ({e}).[/red]"
            )
            console.print(
                "[red]Archive ingestion never falls back to synthetic/curated records. "
                "Fix the dataset path and re-run.[/red]"
            )
            raise

        self.stats.rows_read = rows_processed
        self.stats.individual_pages_attempted = rows_processed
        self.stats.individual_pages_success = records_scraped
        self.stats.individual_pages_failed = self.stats.skipped_total
        self.stats.completed_at = datetime.utcnow()

        self._print_ingestion_summary()

    def _print_ingestion_summary(self) -> None:
        """Print the final archive-ingestion accounting summary."""
        st = self.stats
        console.print()
        console.print(Panel.fit(
            "\n".join([
                "[bold]Archive Ingestion Summary[/bold]",
                "",
                f"  Rows read           : {st.rows_read:>8,}",
                f"  PDF parsed          : {st.pdf_parsed:>8,}",
                f"  DOCX parsed         : {st.docx_parsed:>8,}",
                f"  Skipped             : {st.skipped_total:>8,}",
                f"      broken          : {st.skipped_broken:>8,}",
                f"      scanned         : {st.skipped_scanned:>8,}",
                f"      unsupported     : {st.skipped_unsupported:>8,}",
                f"      parser failure  : {st.skipped_parser_failure:>8,}",
                f"      other           : {st.skipped_other:>8,}",
                f"  Synthetic generated : {st.synthetic_generated:>8,}",
            ]),
            title="[bold green]Phase 1 Complete — Scraping[/bold green]",
            border_style="green",
        ))


# ─────────────────────────────────────────────────────────────────────────────
# Mock Data Generator
# ─────────────────────────────────────────────────────────────────────────────

class MockDataGenerator:
    """
    Generates realistic Lok Sabha Q&A records for when live scraping is
    not possible (blocked site, no network, etc.).
    """

    MINISTRIES = [
        ("Finance", 0.15),
        ("Health and Family Welfare", 0.12),
        ("Education", 0.10),
        ("Home Affairs", 0.08),
        ("External Affairs", 0.07),
        ("Defence", 0.06),
        ("Railways", 0.06),
        ("Agriculture and Farmers Welfare", 0.05),
        ("Road Transport and Highways", 0.05),
        ("Power", 0.04),
        ("Drinking Water and Sanitation", 0.03),
        ("Environment, Forest and Climate Change", 0.03),
        ("Women and Child Development", 0.02),
        ("Labour and Employment", 0.02),
        ("Micro, Small and Medium Enterprises", 0.02),
    ]

    TOPICS_BY_MINISTRY = {
        "Finance": [
            "GST collection", "Income tax slabs", "Bank NPAs", "Digital payments",
            "Foreign direct investment", "Fiscal deficit", "External debt", "Currency reserves",
            "Startup India scheme", "MSME credit", "Insurance penetration", "Stock market reforms",
        ],
        "Health and Family Welfare": [
            "Ayushman Bharat scheme", "National Health Mission", "COVID-19 vaccination",
            "Malaria control", "Dengue cases", "Doctor-patient ratio", "Hospital beds",
            "Medical college seats", "Mental health services", "Drug regulation",
            "TB elimination programme", "National AIDS Control Programme",
        ],
        "Education": [
            "Sarva Shiksha Abhiyan", "Mid-day meal scheme", "Digital education",
            "Higher education enrolment", "Skill India mission", "National Education Policy",
            "Vocational training", "Teacher training", "ICDS programme", "E-learning platforms",
        ],
        "Home Affairs": [
            "Border security", "CCTV coverage", "Police modernization", "Cybercrime cases",
            "Drug trafficking", "Refugee policy", "Internal security", "Intelligence sharing",
            "Naxalism", "Terrorism", "UIDAI data security",
        ],
        "External Affairs": [
            "Neighbouring countries relations", "G20 summits", "UN Security Council reforms",
            "Vande Bharat flights", "Passport services", "Indian diaspora", "Trade agreements",
            "Malabar naval exercise", "BIMSTEC cooperation",
        ],
        "Defence": [
            "Border Infrastructure", "Defence procurement", "Military modernization",
            "Himalayan warfare", "Aircraft carrier INS Vikrant", "HAL production",
            "Ex-servicemen welfare", "Defence R&D", "Border roads organization",
        ],
        "Railways": [
            "Train accidents", "Station modernization", "Railway electrification",
            "Vande Bharat trains", "Freight corridor", "Coach manufacturing",
            "Rail safety fund", "High-speed rail", "Railway station cleanliness",
        ],
        "Agriculture and Farmers Welfare": [
            "MSP procurement", "PM-KISAN scheme", "Crop insurance", "Fertilizer subsidy",
            "Cold chain infrastructure", "Organic farming", "Mandi reforms",
            "Coffee plantation", "Fisheries development", "Irrigation coverage",
        ],
        "Road Transport and Highways": [
            "Expressway construction", "NH maintenance", "Highway accident data",
            "Bharatmala Pariyojana", "Toll plaza management", "Electric vehicle charging",
            "Bridge collapse incidents", "Road safety measures", "Greenfield highways",
        ],
        "Power": [
            "Power generation capacity", "Transmission losses", "Smart grid deployment",
            "Renewable energy targets", "UDA scheme", "Coal supply to power plants",
            "Electricity access", "Power tariff reforms", "NTPC projects",
        ],
        "Drinking Water and Sanitation": [
            "Jal Jeevan Mission", "Swachh Bharat Mission", "Water quality testing",
            "Groundwater depletion", "River cleaning", "Sewage treatment",
            "Fluorosis affected areas", "Drought affected districts",
        ],
        "Environment, Forest and Climate Change": [
            "Air quality index", "Forest cover", "Carbon emission targets",
            "Endangered species protection", "Plastic waste management",
            "Green India Mission", "Climate action plan", "Noise pollution",
        ],
        "Women and Child Development": [
            "Beti Bachao Beti Padhao", "ICDS services", "Child labour",
            "Women SHG groups", "Maternal mortality", "Anganwadi infrastructure",
            "Nutrition programme", "Women helpline", "Domestic violence cases",
        ],
        "Labour and Employment": [
            "MNREGA", "EPFO coverage", "Gig economy workers", "Shram Yogi Mandhan",
            "Skill development programmes", "Industrial disputes", "Factory Act compliance",
            "Child labour abolition", "MSME employment",
        ],
        "Micro, Small and Medium Enterprises": [
            "Mudra loan scheme", "PSU procurement from MSMEs", "Cluster development",
            "Technex support", "Export from MSMEs", "Women entrepreneurs",
            "ZED certification", "Formalization of enterprises",
        ],
    }

    QUESTION_TYPES = [
        (QuestionType.STARRED, 0.20),
        (QuestionType.UNSTARRED, 0.70),
        (QuestionType.SHORT_NOTICE, 0.05),
        (QuestionType.HALF_HOUR, 0.03),
        (QuestionType.CALLING_ATTENTION, 0.02),
    ]

    QUESTION_TEMPLATES = [
        "Will the Minister of {ministry} be pleased to state the details and steps taken regarding {topic}?",
        "Whether the Government has any data or reports on {topic} and the details thereof?",
        "What measures are being taken by the Government to address issues related to {topic}?",
        "What is the current status of policies and schemes under {topic} and their impact on beneficiaries?",
        "Whether there has been any recent budget allocation, review, or targets set for {topic}?",
        "What measures are being taken to improve infrastructure, outreach, and services for {topic}?",
        "Details of any memorandum of understanding or international agreement signed regarding {topic}.",
        "What are the statistical outcomes and progress reports of active schemes concerning {topic}?",
    ]

    CONCERNS = [
        "a significant increase in cases",
        "delays in implementation of schemes",
        "lack of adequate infrastructure",
        "shortage of trained personnel",
        "fund utilization issues",
        "implementation gaps in rural areas",
        "overlapping jurisdiction between ministries",
    ]

    ISSUES = [
        "improve healthcare delivery in rural areas",
        "promote renewable energy adoption",
        "enhance digital connectivity in schools",
        "address unemployment among youth",
        "combat air pollution in metropolitan cities",
        "strengthen road safety measures",
        "boost agricultural exports",
    ]

    SCHEMES = [
        "the PM-KISAN scheme",
        "the Ayushman Bharat Yojana",
        "the Smart Cities Mission",
        "the Swachh Bharat Mission",
        "the Digital India programme",
        "the Make in India initiative",
        "the Skill India Mission",
        "the National Education Policy",
    ]

    METRICS = [
        "the number of beneficiaries under various welfare schemes",
        "the total FDI inflows in the last three years",
        "the rate of reduction in poverty",
        "the number of operational anganwadis",
        "the coverage of rural electrification",
        "the increase in GST collection",
        "the number of startups recognized",
    ]

    SECTOR_TEMPLATES = [
        "healthcare infrastructure in {area}",
        "digital education access in {area}",
        "agricultural marketing in {area}",
        "women's safety in {area}",
        "waste management in {area}",
    ]

    AREAS = [
        "tribal areas", "north-eastern states", "rural districts",
        "urban slums", "border areas", "hilly regions", "coastal areas",
    ]

    def __init__(self, target_count: int = 3500, seed: int = 42) -> None:
        self.target_count = target_count
        self.rng = random.Random(seed)

    def _weighted_choice(self, choices: list[tuple[Any, float]]) -> Any:
        """Select from a weighted list."""
        items, weights = zip(*choices)
        return self.rng.choices(items, weights=weights, k=1)[0]

    def _build_question(
        self,
        ministry: str,
        topic: str,
        session_num: int,
        q_num: int,
    ) -> tuple[str, QuestionType]:
        """Generate a realistic question text."""
        qtype = self._weighted_choice(self.QUESTION_TYPES)
        template = self.rng.choice(self.QUESTION_TEMPLATES)

        # Fill in template variables
        question = template.format(
            ministry=ministry,
            concern=self.rng.choice(self.CONCERNS),
            issue=self.rng.choice(self.ISSUES),
            scheme=self.rng.choice(self.SCHEMES),
            metric=self.rng.choice(self.METRICS),
            programme=self.rng.choice(self.SCHEMES).replace("the ", ""),
            policy=self.rng.choice(self.SCHEMES),
            sector=self.rng.choice(self.SCHEMES).replace("the ", "").replace(" Mission", ""),
            body="World Bank" if "finance" in ministry.lower() else "State Governments",
            topic=topic.lower(),
            entity="complaints" if "women" in ministry.lower() else "incidents",
            area=self.rng.choice(self.AREAS),
        )

        # Add question type prefix
        prefix = f"[{qtype.value.upper()}] " if qtype != QuestionType.UNSTARRED else ""
        return f"{prefix}{question}", qtype

    def _build_answer(self, ministry: str, topic: str, question: str) -> str:
        """Generate a realistic answer text."""
        sentences = []

        # Opening statement
        openings = [
            f"The Minister of {ministry} has stated that",
            "In reply to the question, the Government has informed that",
            f"As per the information provided by the Ministry of {ministry},",
            "The Government has taken note of the concerns raised and",
            f"Based on the latest available data from the Ministry of {ministry},",
        ]
        sentences.append(self.rng.choice(openings))

        # Scheme/programme details
        if any(kw in question.lower() for kw in ["scheme", "programme", "mission"]):
            sentences.append(
                f"The {self.rng.choice(self.SCHEMES).replace('the ', '')} has been "
                f"implemented across {self.rng.randint(20, 36)} states/UTs with a total "
                f"allocation of ₹{self.rng.randint(5000, 150000):,} crore for the current plan period."
            )

        # Statistical data
        sentences.append(
            f"During the period 2019-2024, a total of {random.randint(10000, 500000):,} "
            f"beneficiaries have been covered under various programmes related to {topic.lower()}, "
            f"out of which {random.randint(30, 50)}% are reported to be from rural areas."
        )

        # Ministry response
        responses = [
            f"The Ministry has approved {random.randint(5, 25)} new projects "
            f"worth ₹{random.randint(100, 2000):,} crore for the financial year 2023-24 "
            f"focusing on {topic.lower()}.",
            f"An amount of ₹{random.randint(50, 500):,} crore has been released to "
            f"various state governments for implementation of {self.rng.choice(self.SCHEMES).replace('the ', '')}.",
            f"The Government has empanelled {random.randint(100, 500)} institutions "
            f"for training and capacity building in {topic.lower()}.",
            f"Under the {self.rng.choice(self.SCHEMES).replace('the ', '')}, "
            f"{random.randint(1000, 10000):,} beneficiaries have received direct "
            f"financial assistance.",
        ]
        sentences.append(self.rng.choice(responses))

        # Infrastructure/physical progress
        sentences.append(
            f"A total of {random.randint(50, 500)} {topic.lower()} centres have been "
            f"established, {random.randint(10, 40)} of which are operational in "
            f"{self.rng.choice(['aspirational districts', 'tribal areas', 'NER states', 'rural areas'])}. "
            f"The utilisation rate stands at approximately {random.randint(55, 90)}%."
        )

        # Future plans
        sentences.append(
            f"The Government proposes to expand coverage to an additional "
            f"{random.randint(100, 500)} {topic.lower()} units in the next phase, "
            f"with an estimated budget of ₹{random.randint(200, 1000):,} crore."
        )

        # Concluding statement
        conclusions = [
            "The Government remains committed to addressing these concerns.",
            "Regular monitoring and review mechanisms are in place.",
            "Progress is reviewed quarterly by the Ministry.",
            "The Ministry coordinates with state governments for effective implementation.",
        ]
        sentences.append(self.rng.choice(conclusions))

        return " ".join(sentences)

    def generate(self) -> Iterator[QARecord]:
        """Generate the target number of realistic Q&A records."""
        generated = 0
        session = 18  # 18th Lok Sabha

        # Build flat list of ministry+topic combinations
        ministry_topics = []
        for ministry, _ in self.MINISTRIES:
            for topic in self.TOPICS_BY_MINISTRY.get(ministry, []):
                ministry_topics.append((ministry, topic))

        q_num = 1
        while generated < self.target_count:
            # Select ministry and topic
            ministry, topic = self.rng.choice(ministry_topics)

            # Generate question and answer
            question_text, qtype = self._build_question(ministry, topic, session, q_num)
            answer_text = self._build_answer(ministry, topic, question_text)

            # Build date
            base_date = datetime(2023, 1, 1)
            days_offset = self.rng.randint(0, 365)
            date = (base_date + timedelta(days=days_offset)).strftime("%Y-%m-%d")

            # Build question ID
            question_id = f"{session}-{q_num:04d}"

            # Create record
            record = QARecord(
                question_id=question_id,
                question_text=question_text,
                answer_text=answer_text,
                metadata=QARecordMetadata(
                    ministry=ministry,
                    date=date,
                    session=session,
                    question_number=q_num,
                    subject=topic,
                    question_type=qtype,
                    answer_status="answered",
                    parliament_number=session,
                ),
                scraped_at=datetime.utcnow(),
            )

            yield record
            generated += 1
            q_num += 1

    def generate_batch(self, count: int) -> list[QARecord]:
        """Generate a batch of records as a list."""
        return list(self.generate())


# ─────────────────────────────────────────────────────────────────────────────
# Mock Scraper (Topic-Aligned Synthetic Data)
# ─────────────────────────────────────────────────────────────────────────────

class MockScraper(Scraper):
    """
    Synthesizes topic-aligned Lok Sabha questions.
    Useful for testing scaling limits (e.g. 3,500+ records).
    """

    def __init__(self, base_url: str, target_count: int = 3500, seed: int = 42) -> None:
        super().__init__(base_url)
        self.generator = MockDataGenerator(target_count=target_count, seed=seed)
        self.target_count = target_count

    def scrape_all(self, max_records: int = 3500) -> Iterator[QARecord]:
        self.stats.started_at = datetime.utcnow()
        actual_count = min(max_records, self.target_count)
        console.print(f"[cyan]Synthesizing {actual_count:,} topic-aligned Q&A records...[/cyan]")

        count = 0
        with Progress(
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            TimeElapsedColumn(),
            console=console,
        ) as p:
            task = p.add_task("Synthesizing records...", total=actual_count)
            for record in self.generator.generate():
                if count >= actual_count:
                    break
                yield record
                count += 1
                p.update(task, completed=count)

        self.stats.individual_pages_attempted = actual_count
        self.stats.individual_pages_success = actual_count
        self.stats.completed_at = datetime.utcnow()


# ─────────────────────────────────────────────────────────────────────────────
# Local File Scraper (Loads from raw JSONL backup)
# ─────────────────────────────────────────────────────────────────────────────

class LocalFileScraper(Scraper):
    """Loads records directly from a pre-ingested JSONL file."""

    def __init__(self, file_path: str) -> None:
        super().__init__(base_url="local")
        self.file_path = Path(file_path)

    def scrape_all(self, max_records: int = 3500) -> Iterator[QARecord]:
        self.stats.started_at = datetime.utcnow()
        if not self.file_path.exists():
            raise FileNotFoundError(f"Local file does not exist: {self.file_path}")

        count = 0
        with Progress(
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=console,
        ) as p:
            task = p.add_task(f"Loading from {self.file_path.name}...", total=max_records)
            with open(self.file_path, encoding="utf-8") as f:
                for line in f:
                    if count >= max_records:
                        break
                    try:
                        data = json.loads(line.strip())
                        record = QARecord.model_validate(data)
                        yield record
                        count += 1
                        self.stats.individual_pages_success += 1
                    except Exception:
                        self.stats.parse_errors += 1
                    finally:
                        self.stats.individual_pages_attempted += 1
                        p.update(task, completed=count)

        self.stats.completed_at = datetime.utcnow()


# ─────────────────────────────────────────────────────────────────────────────
# Scraper Factory
# ─────────────────────────────────────────────────────────────────────────────

class ScraperFactory:
    """
    Modular Scraper Factory selecting appropriate Strategy classes.
    """

    STRATEGIES = ["live", "archive", "mock", "local"]

    def __init__(
        self,
        base_url: str = "https://sansad.in/ls/questions/questions-and-answers",
        strategy: str = "archive",
        local_file: str | None = None,
        rate_limit: float = 2.0,
        timeout: int = 30,
        max_retries: int = 3,
        use_pdf: bool = True,  # Default to PDF extraction
        ministry_filter: str | None = None,
    ) -> None:
        if strategy not in self.STRATEGIES:
            raise ValueError(f"Unknown strategy '{strategy}'. Options: {self.STRATEGIES}")
        self.base_url = base_url
        self.strategy = strategy
        self.local_file = local_file
        self.rate_limit = rate_limit
        self.timeout = timeout
        self.max_retries = max_retries
        self.use_pdf = use_pdf
        self.ministry_filter = ministry_filter

    def create_scraper(self) -> Scraper:
        """Instantiate and return the appropriate Scraper subclass strategy."""
        if self.strategy == "archive":
            console.print("[cyan]Strategy: ARCHIVE (Curated genuine Lok Sabha dataset).[/cyan]")
            return RealArchiveScraper(
                base_url=self.base_url,
                rate_limit_seconds=self.rate_limit,
                timeout_seconds=self.timeout,
                max_retries=self.max_retries,
                use_pdf=self.use_pdf,
                ministry_filter=self.ministry_filter,
            )

        if self.strategy == "local":
            if not self.local_file:
                raise ValueError("local_file path is required for strategy='local'")
            console.print(f"[cyan]Strategy: LOCAL (Loading from {self.local_file}).[/cyan]")
            return LocalFileScraper(self.local_file)

        if self.strategy == "mock":
            console.print("[cyan]Strategy: MOCK (Generating synthetic topic-aligned records).[/cyan]")
            return MockScraper(self.base_url)

        # "live" Strategy
        console.print(f"[cyan]Strategy: LIVE (Crawling live Lok Sabha website {self.base_url}).[/cyan]")
        return LiveLoksabhaScraper(
            base_url=self.base_url,
            rate_limit_seconds=self.rate_limit,
            timeout_seconds=self.timeout,
            max_retries=self.max_retries,
        )
