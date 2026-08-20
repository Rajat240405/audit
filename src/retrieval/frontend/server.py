"""
FastAPI Backend Server for Phase 10 Multi-Provider LLM Platform.
Serves static client assets, manages runtime provider/model/mode switches,
and exposes REST API endpoints for Hybrid RAG, GraphRAG, and Graph Explorer lookups.
Features the Provider and Model Registries for intelligent, zero-override model matching.
Refactored to treat Fast/Deep as execution profiles rather than separate model variants.
Supports true runtime dynamic model discovery for Ollama.
Symmetric quality-optimized parameters applied for both standard and thinking models.
Ensures perfect, secure propagation of in-memory API keys across runtime boundaries.
"""

from __future__ import annotations

import io
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import httpx
import uvicorn
from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel

from src.retrieval.hybrid.pipeline import HybridRAGPipeline
from src.retrieval.graph.store import GraphStore
from src.retrieval.graph.retriever import GraphRetriever
from src.generation.client import LLMClient, ollama_base_url
from src.generation.generator import AnswerGenerator
from src.generation.policy import resolve_execution
from src.generation.registry import (
    ModelFamily,
    model_registry,
    provider_registry,
    resolve_family_for_provider,
)
from src.generation.vllm_discovery import (
    VLLMDiscoveryError,
    discover_active_vllm_model,
    refresh_vllm_discovery,
    resolve_served_family,
)
from src.utils.app_paths import (
    corpus_path,
    data_dir,
    ensure_data_dirs,
    graph_dir as resolve_graph_dir,
    inbox_dir,
    index_dir as resolve_index_dir,
    project_root,
    user_knowledge_dir,
)

app = FastAPI(
    title="Parliamentary & Audit Assistant Multi-Provider API",
    description="Provider-agnostic API backing the Interactive Chat Frontend for Phase 10."
)

# Project root + optional APP_DATA_DIR / APP_INDEX_DIR / APP_MODEL_DIR (P0.3).
# Unset env keeps Windows/dev behavior (paths under the repo). Never use CWD.
PROJECT_ROOT = project_root()
ensure_data_dirs()

# ─────────────────────────────────────────────────────────────────────────────
# In-Memory Active Configuration State (Thread-safe single-session storage)
# ─────────────────────────────────────────────────────────────────────────────

from src.generation.defaults import default_family_id, default_model_name, default_num_ctx

# Provider default: env-driven so the SAME code runs on PC (ollama) and
# HPC (vllm). APP_DEFAULT_PROVIDER wins; else VLLM_BASE_URL present -> vllm.
_DEFAULT_PROVIDER = (
    os.environ.get("APP_DEFAULT_PROVIDER", "").strip().lower()
    or ("vllm" if os.environ.get("VLLM_BASE_URL") else "ollama")
)


def _enabled_providers() -> list[str]:
    """Providers THIS deployment exposes — the runtime boundary between PC
    and HPC (no environment-specific ifs anywhere else in the codebase).

    ``APP_PROVIDERS`` (comma-separated) wins; the default is the single
    provider derived from the environment (ollama on PC, vllm when
    VLLM_BASE_URL is set). Unknown names are dropped. Read lazily per call
    so tests and launchers can vary the env without re-importing the app.
    """
    raw = (os.environ.get("APP_PROVIDERS") or "").strip()
    if raw:
        provs = [p.strip().lower() for p in raw.split(",") if p.strip()]
        valid = [p for p in provs if provider_registry.get(p) is not None]
        return valid or [_DEFAULT_PROVIDER]
    return [_DEFAULT_PROVIDER]


_PROVIDER_LABELS = {
    "ollama": "Ollama (local)",
    "vllm": "vLLM (server)",
    "openai_compatible": "OpenAI-compatible (server)",
    "huggingface": "In-process (HuggingFace)",
}


# Boot-time sanity: explicit APP_PROVIDERS that excludes the default provider
# would boot into a provider the deployment says is unavailable — recover to
# the first enabled one loudly rather than failing the first request.
_boot_enabled = _enabled_providers()
if _DEFAULT_PROVIDER not in _boot_enabled:
    print(
        f"[provider] APP_DEFAULT_PROVIDER={_DEFAULT_PROVIDER!r} is not in "
        f"APP_PROVIDERS={_boot_enabled} — booting with {_boot_enabled[0]!r}"
    )
    _DEFAULT_PROVIDER = _boot_enabled[0]

ACTIVE_CONFIG = {
    "provider": _DEFAULT_PROVIDER,
    # Default family/model resolve from src.generation.defaults (single source)
    "model_family": default_family_id(),
    "model": default_model_name(),
    "mode": "fast",            # "fast" or "deep"
    "retrieval_mode": "hybrid",# "hybrid" or "graph"
}


def _active_api_key(provider: str | None = None) -> str | None:
    """In-memory API key slot. Local providers (ollama/huggingface) need none."""
    return None


# Last successfully discovered vLLM serving context (--max-model-len). The
# serving limit feeds effective-context resolution in the policy; when
# discovery is currently failing we keep the last-known value rather than
# silently dropping the clamp (loud note printed).
_LAST_SERVING_LIMIT: int | None = None


def _current_serving_limit(prov: str) -> int | None:
    """Serving context limit for server providers (TTL-cached discovery —
    never a per-request /v1/models call); None for non-server providers and
    when never successfully discovered (the policy then does not clamp on
    the serving input — nothing is invented)."""
    global _LAST_SERVING_LIMIT
    if prov not in ("vllm", "openai_compatible"):
        return None
    try:
        active = discover_active_vllm_model()
    except VLLMDiscoveryError:
        return _LAST_SERVING_LIMIT
    mml = active.get("max_model_len")
    _LAST_SERVING_LIMIT = int(mml) if isinstance(mml, (int, float)) and mml > 0 else None
    return _LAST_SERVING_LIMIT


def _apply_execution_plan(family, prov: str, exec_mode: str):
    """Resolve the execution policy for (family, mode, provider) ONCE and bind
    the plan to the shared client/generator/config state. This is the ONLY
    mode→parameters path in the server (previously copy-pasted per handler).

    Reasoning models (e.g. Qwen3.6) consume max_tokens with their
    chain-of-thought BEFORE writing the answer — the Deep plan carries the
    reasoning reserve so the answer budget survives (see policy.py).

    Context: the policy receives the currently-served vLLM limit (when
    discoverable) so the effective context is
    min(native ∩ serving ∩ app ceiling) — never a universal constant.
    """
    plan = resolve_execution(
        family, exec_mode, prov, serving_limit=_current_serving_limit(prov)
    )
    resolved_model = family.model_name

    llm_client.provider = prov
    llm_client.model = resolved_model
    llm_client.num_ctx = plan.num_ctx
    llm_client.temperature = plan.temperature
    llm_client.max_tokens = plan.max_tokens
    # Wire-level think from the plan: equals plan.thinking for catalogued
    # models (capability known); None for dynamically discovered unknown
    # models — the provider adapters then send NO thinking control at all
    # (their `think is not None` guards), so unknown models run at server
    # default and never receive an invented control. plan.thinking stays the
    # mode request for reasoning-display gating.
    llm_client.think = plan.wire_think
    ACTIVE_CONFIG["verify_depth"] = plan.verify_depth
    ACTIVE_CONFIG["model"] = resolved_model
    # Task 3: attach the resolved plan — the generator switches from legacy
    # fixed-count/char-caps to reserve-based dynamic evidence budgeting.
    generator.plan = plan
    # Legacy fallback knobs still bound (used only if the plan were detached)
    generator.max_doc_chars = plan.max_doc_chars
    generator.max_context_docs = plan.max_context_docs
    for w in plan.warnings:
        print(f"[exec] warning: {w}")
    return plan, resolved_model


def _refresh_active_served_family() -> None:
    """Re-resolve the ACTIVE served model for vLLM / OpenAI-compatible
    providers before a generation request. TTL-cached (no /v1/models call per
    request); an explicit VLLM_MODEL pin freezes the choice (override
    semantics); non-server providers are a no-op. If the served model changed
    since boot (HPC swap), ACTIVE_CONFIG follows the server — the model id
    sent on the wire is always the currently-served one. Discovery failures
    keep the last-known-good configuration (loudly)."""
    prov = ACTIVE_CONFIG["provider"]
    if prov not in ("vllm", "openai_compatible"):
        return
    if (os.environ.get("VLLM_MODEL") or "").strip():
        return  # explicit pin wins; nothing to discover for the wire name
    try:
        active = discover_active_vllm_model()
    except VLLMDiscoveryError as e:
        print(f"[models] discovery refresh failed ({e}) — keeping {ACTIVE_CONFIG['model']}")
        return
    if active["id"] == ACTIVE_CONFIG["model"]:
        return
    fam, source = resolve_served_family(
        active["id"], active.get("max_model_len"), provider=prov
    )
    print(f"[models] served model changed: {ACTIVE_CONFIG['model']} -> "
          f"{active['id']} (metadata={source})")
    ACTIVE_CONFIG["model_family"] = fam.id
    ACTIVE_CONFIG["model"] = fam.model_name


# ── Task 3 dynamic evidence budgeting — server-side helpers ────────────────
def _effective_top_k(request_top_k, plan) -> int:
    """Retrieval candidate pool for this request. The profile's initial pool
    (fast 5 / deep 10 — an INITIAL pool, not a final document quota) is the
    floor; an explicitly larger request top_k is still honored (API compat).
    Budget-driven admission downstream decides what actually reaches the prompt.
    """
    pool = getattr(plan, "retrieval_top_k", 5) or 5
    try:
        req = int(request_top_k or 0)
    except (TypeError, ValueError):
        req = 0
    return max(req, pool, 1)


def _admission_diag(question: str, results) -> tuple[list[str], dict]:
    """Admitted ids + full admission diagnostics for (question, results).

    Planned path (ExecutionPlan attached): budget-driven admission via the
    generator's shared prepare_context cache — the SAME assembly generate()/
    generate_stream() then reuse (stream/non-stream parity). The diagnostics
    carry the fields needed to distinguish, at runtime:
      CASE A  pool > admitted  → correct Task-3 budget admission (skipped ids listed)
      CASE B  retrieved < retrieval_top_k → narrowed upstream of admission
              (see the per-stage counts in the server trace line); NOT a
              top_k-propagation issue when effective_top_k in the trace is the
              plan pool (5 Standard / 10 Deep)
      CASE C  legacy_max_context_docs_fallback=True → plan was MISSING on a
              request that should be planned (alarm — planned requests must
              never hit the legacy cap)
    """
    plan = getattr(generator, "plan", None)
    if plan is None:
        cap = max(1, int(getattr(generator, "max_context_docs", 5) or 5))
        ids = [r.doc_id for r in results[:cap]]
        return ids, {
            "plan_attached": False,
            "legacy_max_context_docs_fallback": True,
            "legacy_cap": cap,
            "pool": len(results),
            "admitted": len(ids),
            "skipped_doc_ids": [r.doc_id for r in results[cap:]],
            "evidence_budget_tokens": None,
            "evidence_used_tokens": None,
        }
    _, ids, diag = generator.prepare_context(question, results)
    diag["plan_attached"] = True
    diag["legacy_max_context_docs_fallback"] = False
    diag["retrieved_count"] = len(results)   # hybrids-stage parents (pre-slice)
    diag["retrieval_top_k"] = getattr(plan, "retrieval_top_k", None)
    diag["plan_evidence_budget_tokens"] = getattr(plan, "evidence_budget_tokens", None)
    return ids, diag


def _admitted_ids(question: str, results) -> list[str]:
    """Budget-admitted doc ids for (question, results) via the generator's
    shared prepare_context cache — the SAME assembly generate()/
    generate_stream() then reuse, so the sources the UI shows are exactly the
    evidence the model receives (stream/non-stream parity)."""
    ids, _ = _admission_diag(question, results)
    return ids


def _filter_to_admitted(sources, admitted_ids, key=None) -> list:
    keep = set(admitted_ids)
    if key is None:
        key = lambda s: getattr(s, "doc_id", None)  # noqa: E731
    return [s for s in sources if key(s) in keep]


def _maybe_enrich_deep_neighbors(plan, results) -> None:
    """Deep-only neighbor-chunk pull-in: re-bond heading-like previous chunks
    to mid-document long-chunk evidence (uses metadata already persisted in
    the index; no reindexing). Best-effort — retrieval must never fail on it."""
    if getattr(plan, "mode", "") != "deep":
        return
    try:
        from src.generation.evidence import enrich_deep_neighbors

        n = enrich_deep_neighbors(results, getattr(pipeline, "_long_chunk_map", {}) or {})
        if n:
            print(f"[context] deep neighbor pull-in: {n} heading chunk(s) bonded")
    except Exception as e:  # noqa: BLE001
        print(f"[context] neighbor pull-in skipped ({type(e).__name__}: {e})")

# ─────────────────────────────────────────────────────────────────────────────
# Lazy pipeline loading
# ─────────────────────────────────────────────────────────────────────────────
# The Hybrid RAG pipeline loads the bge-m3 embedding model (weights) AND the
# FAISS/BM25 index at construction. Loading it at import time delays the
# server binding its port by many seconds — the frontend sees ECONNREFUSED
# on /api/status meanwhile. Instead we expose a lazy proxy: the port binds
# immediately, and the heavy model/index load happens once, on the first
# retrieval request (a one-time delay on the first query only).
import threading as _threading


class _LazyPipeline:
    """Builds the real HybridRAGPipeline on first use; forwards all access."""

    def __init__(self) -> None:
        self._instance = None
        self._lock = _threading.Lock()

    def _get(self):
        if self._instance is None:
            with self._lock:
                if self._instance is None:
                    self._instance = self._build()
        return self._instance

    def swap(self, new_instance) -> None:
        """Atomically replace the live pipeline under the same lock used by
        _get(), so an ingest rebuild can never race a query's load."""
        with self._lock:
            self._instance = new_instance

    @staticmethod
    def _build():
        p = HybridRAGPipeline()
        index_dir = resolve_index_dir()
        if index_dir.exists():
            p.load(index_dir)
            chunks_count = (
                len(p._chunk_map) if p.use_chunking else len(p._doc_map)
            )
            print("=" * 60)
            print(f"Embedding Model : {p.embedder.model_name}")
            print(f"Embedding Dimension : {p.embedder.embedding_dim}")
            print(f"Documents Indexed : {len(p._doc_map)}")
            print(f"Chunks Indexed : {chunks_count}")
            print("FAISS Index Built Successfully")
            print("=" * 60)
        return p

    def __getattr__(self, name):
        return getattr(self._get(), name)


index_dir = resolve_index_dir()
graph_dir = resolve_graph_dir()

# Lazy: server binds port immediately; model+index load on first query.
pipeline = _LazyPipeline()

graph_store = GraphStore(storage_dir=str(graph_dir))
if graph_store.graph_file.exists():
    graph_store.load()
graph_retriever = GraphRetriever(store=graph_store)

# Resolve default starting configuration dynamically from registry. When the
# env selects a non-Ollama provider (HPC: APP_DEFAULT_PROVIDER=vllm), the
# global default family (PC/ollama "qwen3") is not served by that provider —
# re-resolve to a family the provider actually serves so the first request
# cannot 404 against vLLM. PC/ollama boot is unchanged.
#
# vLLM/openai_compatible boot order (DISCOVERY-AWARE, VLLM_MODEL optional):
#   1. explicit VLLM_MODEL pin      → exact catalog resolution, or a dynamic
#                                     conservative family when uncatalogued;
#   2. otherwise /v1/models         → the ACTIVE served model (deterministic
#                                     policy, see vllm_discovery), resolved
#                                     against the catalog by EXACT name;
#   3. discovery failure            → loud warning + the legacy first-family
#                                     fallback (the app still boots while the
#                                     server is down; per-request refresh
#                                     retries discovery before generation).
def _boot_served_family(provider: str) -> Optional[ModelFamily]:
    if provider not in ("vllm", "openai_compatible"):
        return None
    try:
        active = discover_active_vllm_model()  # honors the VLLM_MODEL pin itself
    except VLLMDiscoveryError as e:
        print(f"[models] served-model discovery unavailable at boot ({e}) — "
              "falling back to catalog selection; will retry before generation")
        return resolve_family_for_provider(
            model_registry, provider, ACTIVE_CONFIG["model_family"],
            preferred_model=os.environ.get("VLLM_MODEL"),
        )
    fam, source = resolve_served_family(
        active["id"], active.get("max_model_len"), provider=provider
    )
    print(f"[models] active served model: {active['id']} "
          f"(metadata={source}, pinned={active['pinned']}"
          + (f", alternatives={active['alternatives']}" if active["alternatives"] else "")
          + ")")
    return fam


_default_fam = _boot_served_family(ACTIVE_CONFIG["provider"]) or resolve_family_for_provider(
    model_registry,
    ACTIVE_CONFIG["provider"],
    ACTIVE_CONFIG["model_family"],
    preferred_model=os.environ.get("VLLM_MODEL"),
)
if _default_fam:
    ACTIVE_CONFIG["model_family"] = _default_fam.id
    ACTIVE_CONFIG["model"] = _default_fam.model_name
    _num_ctx = _default_fam.context_window
else:
    _num_ctx = default_num_ctx()

# Instantiate the active client and generator router
llm_client = LLMClient(
    provider=ACTIVE_CONFIG["provider"],
    model=ACTIVE_CONFIG["model"],
    num_ctx=_num_ctx
)
generator = AnswerGenerator(llm_client=llm_client)

# Dedicated small/fast model for AI-edit buttons (bullets/formal/concise/
# grammar/prose). Simple rewrites don't need the big model — a small local
# model makes edits finish in seconds instead of a minute.
# Override with GRAPHRAG_EDIT_MODEL env (e.g. "qwen2.5:0.5b" for fastest).
_EDIT_MODEL = os.environ.get("GRAPHRAG_EDIT_MODEL", "qwen2.5:3b")
_edit_ctx = default_num_ctx()
try:
    _edit_fam = model_registry.get(_EDIT_MODEL)
    if _edit_fam:
        _edit_ctx = _edit_fam.context_window
except Exception:
    _edit_ctx = default_num_ctx()
edit_llm_client = LLMClient(
    provider=ACTIVE_CONFIG["provider"],
    model=_EDIT_MODEL,
    num_ctx=_edit_ctx,
    max_tokens=2048,
)

# ─────────────────────────────────────────────────────────────────────────────
# Request / Response Schemas
# ─────────────────────────────────────────────────────────────────────────────

class ProviderSwitchRequest(BaseModel):
    provider: str
    model: str # Now serves as the model family ID (e.g. "qwen3")
    api_key: Optional[str] = None


class ChatStreamRequest(BaseModel):
    message: str
    mode: str = "fast"            # Execution Mode: "fast" or "deep"
    retrieval_mode: str = "hybrid"  # Retrieval Mode: "hybrid" or "graph"
    top_k: int = 5
    draft_style: Optional[str] = None  # e.g. formal / concise / executive
    doc_types: Optional[list[str]] = None  # source filter: parliament / annual_report / ...
    orgs: Optional[list[str]] = None          # source filter: expanded org slugs (tree rule already applied)
    doc_categories: Optional[list[str]] = None  # source filter: annual / monthly / budget / ...


class ChatRequest(ChatStreamRequest):
    """Non-streaming chat request — shares the streaming schema (P1.1).

    The /api/chat endpoint reads top_k / doc_types / orgs / doc_categories;
    previously ChatRequest lacked them and every non-streaming call raised
    AttributeError. Frontend uses /api/chat/stream, which is why this stayed
    hidden."""

    pass

class SourceItem(BaseModel):
    doc_id: str
    ministry: Optional[str] = None
    subject: Optional[str] = None
    date: Optional[str] = None  # R1: raw record date stamp, when the record carries one
    document_type: Optional[str] = None
    score: float
    question: str
    answer: str
    # Retrieval trace — per-component scores (hybrid path only; None for graph path)
    dense_score: Optional[float] = None
    bm25_score: Optional[float] = None
    rrf_score: Optional[float] = None
    rerank_score: Optional[float] = None


class ChatResponse(BaseModel):
    answer: str
    sources: List[SourceItem]
    retrieval_latency_ms: float
    generation_latency_ms: float
    prompt_tokens: int
    completion_tokens: int
    total_tokens: int
    is_fallback: bool = False
    is_graph_result: bool = False
    active_provider: str
    active_model: str # Concrete model name
    active_mode: str
    # Expanded runtime metrics (Objective 7)
    model_family: str
    resolved_model: str
    context_window: int
    prompt_budget: int
    network_latency_ms: Optional[float] = None
    # Retrieval trace: per-component scores per hit + stage timings (ms)
    trace: Optional[dict] = None


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def is_semantic_synthesis_query(query: str) -> bool:
    """
    Detects if the query is requesting summarization, explanation, comparison,
    reasoning, synthesis, or drafting over the documents (requiring LLM generation).
    Uses strict whole-word set intersection to prevent substring collisions.
    """
    import re
    query_clean = re.sub(r"[^\w\s]", " ", query.lower()).strip()
    words = set(query_clean.split())
    
    synthesis_triggers = {
        "summarize", "summarise", "explain", "compare", "reason", "why", "synthesize",
        "synthesise", "how", "relation", "contrast", "summary", "analyze", "analyse",
        "draft"
    }
    return bool(words.intersection(synthesis_triggers))


def is_metadata_query(query: str) -> bool:
    """
    Detects if the query is requesting document listings, ministry lookups,
    subject lookups, session/date filtering, graph stats, or other metadata lookups.
    """
    import re
    query_clean = re.sub(r"[^\w\s]", " ", query.lower()).strip()
    words = set(query_clean.split())
    
    metadata_triggers = {
        "list", "find", "lookup", "show", "get", "statistics", "stats", "mps", 
        "members", "session", "date", "document", "question", "questions"
    }
    
    # Also check multi-word phrase triggers
    has_phrase = any(phrase in query_clean for phrase in ["who asked", "which ministry", "what subjects", "where is"])
    
    has_metadata_trigger = bool(words.intersection(metadata_triggers)) or has_phrase
    has_synthesis_trigger = is_semantic_synthesis_query(query)
    
    return has_metadata_trigger and not has_synthesis_trigger


def extract_informative_summary(answer_text: str, subject_text: str) -> str:
    """
    Cleans up official parliamentary answers by removing boilerplate headings,
    minister signatures, salutations, and returns the first 1-2 informative sentences.
    Falls back to the document's subject if no informative text remains.
    """
    import re
    # 1. Standardize whitespace and remove newlines for clean regex matching
    text = re.sub(r'\s+', ' ', answer_text).strip()
    
    # 2. Remove ANSWER/REPLY heading
    text = re.sub(r'(?i)^ANSWER\s*[:\-]?\s*', '', text).strip()
    text = re.sub(r'(?i)^REPLY\s*[:\-]?\s*', '', text).strip()
    
    # 3. Remove MINISTER signatures and departments
    text = re.sub(r'(?i)^MINISTER\s+OF\s+STATE\s*\([^)]*\)\s*(?:IN\s+THE\s+MINISTRY\s+OF\s+[\w\s&,]+)?\s*', '', text).strip()
    text = re.sub(r'(?i)^(?:THE\s+HON’BLE\s+)?MINISTER\s+OF\s+[\w\s&,]+\s*\([^)]*\)\s*', '', text).strip()
    text = re.sub(r'(?i)^(?:THE\s+HON’BLE\s+)?MINISTER\s+OF\s+[\w\s&,]+\s*', '', text).strip()
    
    # 4. Remove leading list indices like (a) to (c) or (a) & (b) or (a)
    text = re.sub(r'(?i)^\([a-g]\)\s*(?:to|&|and)?\s*(?:\([a-g]\))?\s*[:\-\.]?\s*', '', text).strip()
    
    # 5. Remove any residual leading punctuation/dashes
    text = re.sub(r'^[\s\-\:\.\*]+', '', text).strip()

    # Split into sentences
    sentences = re.split(r'\.\s+(?=[A-Z])', text)
    sentences = [s.strip() for s in sentences if s.strip()]

    informative_sentences = []
    for s in sentences:
        s_lower = s.lower()
        if len(s) < 15:
            continue
        if any(term in s_lower for term in ['referred to', 'is enclosed', 'laid on the table']):
            continue
        informative_sentences.append(s)

    if informative_sentences:
        summary = '. '.join(informative_sentences[:2])
        if not summary.endswith('.'):
            summary += '.'
        if len(summary) > 250:
            summary = summary[:247] + '...'
        return summary

    return subject_text if subject_text else 'No detailed summary available.'


# ─────────────────────────────────────────────────────────────────────────────
# REST Endpoints
# ─────────────────────────────────────────────────────────────────────────────

_SPA_RESERVED_PREFIXES = ("api", "health", "assets")


def _serve_spa_index() -> HTMLResponse:
    """Vite index.html (or legacy template). Used for / and client-route refresh."""
    dist = PROJECT_ROOT / "frontend" / "dist"
    dist_index = dist / "index.html"
    if dist_index.exists():
        return HTMLResponse(dist_index.read_text(encoding="utf-8"))
    html_path = Path(__file__).parent / "index.html"
    if not html_path.exists():
        raise HTTPException(status_code=404, detail="index.html template not found")
    return HTMLResponse(html_path.read_text(encoding="utf-8"))


@app.get("/", response_class=HTMLResponse)
async def get_index():
    """Serve the React workstation (production build) if present, else the
    legacy single-file HTML client. Lets `npm run build` output be served by
    the FastAPI backend at the same origin (no CORS, no separate host)."""
    return _serve_spa_index()


@app.get("/assets/{path:path}")
async def get_assets(path: str):
    """Serve built frontend assets (JS/CSS) from frontend/dist/assets."""
    dist = PROJECT_ROOT / "frontend" / "dist"
    asset = dist / "assets" / path
    if not asset.exists():
        raise HTTPException(status_code=404, detail="asset not found")
    media = "application/javascript" if asset.suffix == ".js" else (
        "text/css" if asset.suffix == ".css" else "application/octet-stream"
    )
    return Response(content=asset.read_bytes(), media_type=media)


# ─────────────────────────────────────────────────────────────────────────────
# Served-model discovery
#
# The YAML catalog (config/models.yaml) is METADATA ONLY — it is never proof
# that a model is installed. Availability comes exclusively from the provider
# the app is connected to:
#   Ollama              -> GET  /api/tags  (+ POST /api/show for context_length)
#   vLLM / OpenAI-compat-> GET /v1/models  (incl. server-reported max_model_len)
# A model file on someone's disk that the connected server does NOT serve is
# invisible to Audit; a model served by that server is usable even when it is
# absent from the catalog (registered dynamically with honestly-sourced or
# explicitly-flagged fallback metadata).
# ─────────────────────────────────────────────────────────────────────────────

# Conservative last-resort context window (8192) for models reporting none:
# vllm_discovery.default_unknown_context() is the single definition, always
# flagged metadata_source="fallback" so the UI can show it as assumed (and it
# errs small for budget safety).


def _family_entry(f: ModelFamily, metadata_source: str) -> dict:
    return {
        "id": f.id,
        "display_name": f.display_name,
        "provider": f.provider,
        "model_name": f.model_name,
        "context_window": f.context_window,
        "thinking_capable": f.thinking_capable,
        # Tri-state capability (additive; true/false known, null = unknown).
        # Unknown models are never claimed thinking-capable — and nothing
        # thinking-related is sent to them on the wire (plan.wire_think=None).
        "thinking_supported": (f.thinking.supported if f.thinking else None),
        "recommended_execution_mode": f.recommended_execution_mode,
        "think_mode": f.think_mode,
        "served": True,
        "metadata_source": metadata_source,
    }


def _served_family_entry(prov: str, served_id: str, server_ctx: int | None) -> dict:
    """Map one served model id to a family entry. Metadata resolution order:
    1. exact match in THIS provider's catalog section -> catalog metadata;
    2. exact model_name match in another provider's section -> metadata
       reused, transport (think_mode) follows the active provider;
    3. server-reported context (vLLM max_model_len / Ollama /api/show);
    4. flagged conservative fallback (never presented as detected).

    Thin wrapper over vllm_discovery.resolve_served_family — ONE
    implementation for the UI endpoint and the generation path, so both
    resolve identically.
    """
    fam, source = resolve_served_family(
        served_id,
        server_ctx if isinstance(server_ctx, int) and server_ctx > 0 else None,
        provider=prov,
    )
    return _family_entry(fam, source)


@app.get("/api/providers")
async def get_providers():
    """Providers enabled in THIS deployment (APP_PROVIDERS). The inactive
    provider is not offered: the UI renders this list verbatim and
    /api/provider + /api/models reject anything outside it."""
    active = ACTIVE_CONFIG["provider"]
    return [
        {
            "name": p,
            "label": _PROVIDER_LABELS.get(p, p),
            "active": p == active,
        }
        for p in _enabled_providers()
    ]


@app.get("/api/models")
async def get_models(provider: str):
    """Models available for the selected provider, DISCOVERED from the
    connected server — never from the YAML catalog alone."""
    prov = provider.lower().strip()

    if prov not in _enabled_providers():
        raise HTTPException(
            status_code=403,
            detail=f"Provider {prov!r} is not enabled in this deployment.",
        )

    # ── Ollama: models installed in the connected Ollama service ────────
    if prov == "ollama":
        ollama = provider_registry.get("ollama")
        try:
            tags = ollama.list_tags(base_url=ollama_base_url())
        except Exception:
            raise HTTPException(
                status_code=503,
                detail=f"Ollama service is offline or unreachable at {ollama_base_url()}",
            )
        return [
            _served_family_entry(prov, tag, ollama.show_context_length(tag))
            for tag in tags
        ]

    # ── vLLM / any OpenAI-compatible server: what the server SERVES ─────
    if prov in ("vllm", "openai_compatible"):
        compat = provider_registry.get(prov)
        try:
            served = compat.served_models()
        except Exception:
            raise HTTPException(
                status_code=503,
                detail=f"{prov} server is offline or unreachable at {compat.base_url}",
            )
        return [
            _served_family_entry(prov, s["id"], s.get("max_model_len"))
            for s in served
        ]

    # ── In-process provider: the catalog IS the install list ────────────
    if prov == "huggingface":
        return [
            _family_entry(f, "catalog")
            for f in model_registry.list_by_provider(prov)
        ]

    raise HTTPException(status_code=400, detail="Unknown provider requested")


@app.post("/api/provider")
async def switch_provider(request: ProviderSwitchRequest):
    """
    Switches active provider and model family in memory without restarting FastAPI.
    Also registers the Groq API key in-memory for the current session.
    """
    prov = request.provider.lower().strip()
    # Backend enforcement of the deployment boundary (not just UI hiding):
    # the provider must be enabled for THIS deployment via APP_PROVIDERS.
    if prov not in _enabled_providers():
        raise HTTPException(
            status_code=403,
            detail=f"Provider {prov!r} is not enabled in this deployment.",
        )

    family_id = request.model
    family = model_registry.get(family_id)
    if not family:
        raise HTTPException(status_code=400, detail="Model family not found in registry")

    ACTIVE_CONFIG["provider"] = prov
    ACTIVE_CONFIG["model_family"] = family.id

    # Resolve concrete model parameters through the execution policy
    exec_mode = ACTIVE_CONFIG["mode"]
    plan, resolved_model = _apply_execution_plan(family, prov, exec_mode)

    llm_client.api_key = _active_api_key(prov)  # session API key (slot or env)
    generator.llm_client = llm_client

    return {
        "status": "success",
        "active_provider": prov,
        "active_model_family": family.display_name,
        "resolved_model": resolved_model,
        "context_window": plan.num_ctx,
        "prompt_budget": plan.prompt_budget_tokens,
        "thinking_capable": family.thinking_capable,
        "recommended_execution_mode": family.recommended_execution_mode,
        "is_connected": llm_client.check_health(api_key=_active_api_key(prov))
    }


@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    """
    API routing chat requests dynamically based on execution mode and pathways.
    """
    query = request.message.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Query message cannot be empty")

    ret_mode = request.retrieval_mode.lower()
    exec_mode = request.mode.lower() # "fast" or "deep"
    ACTIVE_CONFIG["mode"] = exec_mode

    # Follow the currently-served model for server providers (TTL-cached —
    # no /v1/models call per request; no-op for Ollama/HF + VLLM_MODEL pins)
    _refresh_active_served_family()

    # Resolve Model Family dynamically from Registry
    family_id = ACTIVE_CONFIG["model_family"]
    family = model_registry.get(family_id) or model_registry.get("qwen2.5")

    # Resolve execution parameters through the execution policy (single source)
    plan, resolved_model = _apply_execution_plan(family, ACTIVE_CONFIG["provider"], exec_mode)
    llm_client.api_key = _active_api_key()  # Propagate cached API key dynamically!

    # Task 3: the reported prompt budget is now the REAL, reserve-based
    # evidence budget the allocator enforces (was the informational-only
    # 0.80 × num_ctx figure).
    prompt_budget = plan.evidence_budget_tokens
    generator.context_budget_ratio = 0.80

    t_ret_start = time.perf_counter()
    sources: List[SourceItem] = []

    # ── Path Selection Matrix ──
    is_graph_result = (ret_mode == "graph" and not is_semantic_synthesis_query(query))

    if is_graph_result:
        # ── PATH A: DETERMINISTIC METADATA QUERY PATH ──
        results = graph_retriever.retrieve(query, top_k=5)
        ret_latency = (time.perf_counter() - t_ret_start) * 1000
        
        # Format Retrieved Results into Source Items
        for r in results:
            sources.append(SourceItem(
                doc_id=r.doc_id,
                ministry=r.metadata.get("ministry") or "-",
                subject=r.metadata.get("subject") or "-",
                score=float(r.score),
                question=r.question,
                answer=r.answer
            ))

        if results:
            header = "### 🕸️ GraphRAG: Document Explorer\n"
            header += "The following real parliamentary document cards were resolved directly from the metadata graph relationships:\n\n"
            
            cards = []
            for r in results:
                summary = extract_informative_summary(r.answer, r.metadata.get("subject", ""))
                
                card = (
                    f"#### 📄 **Document: {r.doc_id}**\n"
                    f"* 🏢 **Ministry**: {r.metadata.get('ministry') or '-'}\n"
                    f"* 🏷️ **Subject**: {r.metadata.get('subject') or '-'}\n"
                    f"* 📅 **Date**: {r.metadata.get('date') or '-'}\n"
                    f"* ❓ **Question Type**: {r.metadata.get('question_type') or 'Unstarred'}\n"
                    f"* 💡 **Summary**: {summary}\n"
                )
                cards.append(card)
            
            answer = header + "\n---\n".join(cards)
        else:
            answer = (
                "### 🕸️ GraphRAG: Document Explorer\n\n"
                "No matching parliamentary metadata relationships were resolved for your query in the graph."
            )
            
        comp_tok = len(answer) // 4

        return ChatResponse(
            answer=answer,
            sources=sources,
            retrieval_latency_ms=ret_latency,
            generation_latency_ms=0.0,
            prompt_tokens=0,
            completion_tokens=comp_tok,
            total_tokens=comp_tok,
            is_fallback=False,
            is_graph_result=True,
            active_provider=ACTIVE_CONFIG["provider"],
            active_model=resolved_model,
            active_mode=exec_mode,
            model_family=family.display_name,
            resolved_model=resolved_model,
            context_window=family.context_window,
            prompt_budget=prompt_budget,
            network_latency_ms=0.0
        )

    else:
        # ── PATH B: SEMANTIC / SYNTHESIS QUERY PATH ──
        if ret_mode == "graph":
            results = graph_retriever.retrieve(query, top_k=request.top_k)
            ret_latency = (time.perf_counter() - t_ret_start) * 1000
        else:
            results, timings = pipeline.retrieve(
                query, top_k=_effective_top_k(request.top_k, plan),
                doc_types=request.doc_types,
                orgs=request.orgs, doc_categories=request.doc_categories,
            )
            ret_latency = (time.perf_counter() - t_ret_start) * 1000
            # Task 3 (Deep): re-bond heading-like neighbor chunks before
            # evidence is budgeted — metadata already in the index maps.
            _maybe_enrich_deep_neighbors(plan, results)

        # Format Retrieved Results into Source Items
        for r in results:
            sources.append(SourceItem(
                doc_id=r.doc_id,
                ministry=r.metadata.get("ministry") or "-",
                subject=r.metadata.get("subject") or "-",
                date=r.metadata.get("date"),
                score=float(r.score),
                question=r.question,
                answer=r.answer,
                dense_score=float(r.dense_score) if r.dense_score is not None else None,
                bm25_score=float(r.bm25_score) if r.bm25_score is not None else None,
                rrf_score=float(r.rrf_score) if r.rrf_score is not None else None,
                rerank_score=float(r.rerank_score) if r.rerank_score is not None else None,
            ))

        # Task 3: the model receives the budget-admitted set (not a fixed
        # count). prepare_context runs ONCE here (pure, no LLM); generate()
        # reuses the cached assembly — "what the model received" == sources.
        if ret_mode != "graph" and getattr(generator, "plan", None) is not None:
            admitted = _admitted_ids(query, results)
            if len(sources) != len(admitted):
                print(
                    f"[context] pool={len(sources)} admitted={len(admitted)} "
                    "(budget-driven)"
                )
            sources = _filter_to_admitted(sources, admitted)

        t_gen_start = time.perf_counter()
        
        api_key = _active_api_key()
        llm_available = llm_client.check_health(api_key=api_key)

        network_latency_ms = 0.0
        gen_error: Optional[str] = None

        if llm_available and results:
            try:
                gen_res = generator.generate(query, results)
                gen_latency = gen_res.generation_latency_ms
                answer = gen_res.answer
                prompt_tok = gen_res.prompt_tokens
                comp_tok = gen_res.completion_tokens
                total_tok = gen_res.total_tokens
                is_fallback = False
            except Exception as e:  # noqa: BLE001 - generation must never 500 the endpoint
                # An unhandled generation error (Ollama context-length exceeded,
                # timeout, connection drop, non-413 HTTP error) currently
                # surfaces as HTTP 500 with no trace. Log the REAL cause and
                # return the retrieved sources with a graceful notice instead.
                import traceback
                print(f"[Generation failed] {type(e).__name__}: {e}")
                print(traceback.format_exc(limit=5))
                gen_error = f"{type(e).__name__}: {str(e)[:300]}"
                is_fallback = True
                gen_latency = (time.perf_counter() - t_gen_start) * 1000
                provider_label = {
                    "ollama": "Ollama (Local)",
                    "huggingface": "HuggingFace (In-container)",
                }.get(ACTIVE_CONFIG["provider"], ACTIVE_CONFIG["provider"].title())
                answer = (
                    f"**[System Notice: {provider_label} Generation Failed]**\n\n"
                    f"The retrieved documents below were found, but the LLM could "
                    f"not generate an answer for this query.\n"
                    f"* **Error**: `{gen_error}`\n\n"
                    f"**Suggestions**:\n"
                    f"1. Switch to **Fast Mode** (lighter context budget) and retry.\n"
                    f"2. Make the question more specific.\n"
                    f"3. If this recurs, check the server console for the full traceback.\n"
                )
                prompt_tok = 0
                comp_tok = len(answer) // 4
                total_tok = comp_tok
        else:
            is_fallback = True
            gen_latency = (time.perf_counter() - t_gen_start) * 1000
            provider_label = {
                "ollama": "Ollama (Local)",
                "huggingface": "HuggingFace (In-container)",
            }.get(ACTIVE_CONFIG["provider"], ACTIVE_CONFIG["provider"].title())
            err_cause = "Ollama Offline: The local service is currently offline or unreachable."

            answer = (
                f"**[System Notice: {provider_label} Generation Offline]**\n\n"
                f"The active LLM service is currently offline or unreachable.\n"
                f"* **Reason**: {err_cause}\n\n"
                f"Because this query requires cognitive synthesis, explanation, or comparison, a complete "
                f"response cannot be compiled. Please start the local service."
            )
            prompt_tok = 0
            comp_tok = len(answer) // 4
            total_tok = comp_tok

        trace_payload = None
        if ret_mode != "graph":
            trace_payload = {
                "dense_search_ms": round(timings.dense_search_ms, 2),
                "bm25_search_ms": round(timings.bm25_search_ms, 2),
                "rrf_fusion_ms": round(timings.rrf_fusion_ms, 2),
                "rerank_ms": round(timings.rerank_ms, 2),
                "embed_query_ms": round(timings.embed_query_ms, 2),
                "retrieval_total_ms": round(ret_latency, 2),
            }

        return ChatResponse(
            answer=answer,
            sources=sources,
            retrieval_latency_ms=ret_latency,
            generation_latency_ms=gen_latency,
            prompt_tokens=prompt_tok,
            completion_tokens=comp_tok,
            total_tokens=total_tok,
            is_fallback=is_fallback,
            is_graph_result=is_graph_result,
            active_provider=ACTIVE_CONFIG["provider"],
            active_model=resolved_model,
            active_mode=exec_mode,
            model_family=family.display_name,
            resolved_model=resolved_model,
            context_window=family.context_window,
            prompt_budget=prompt_budget,
            network_latency_ms=network_latency_ms,
            trace=trace_payload
        )


# ─────────────────────────────────────────────────────────────────────────────
# Workstation API (frontend redesign: streaming, status, build, export)
# ─────────────────────────────────────────────────────────────────────────────

def _to_sources(results: list) -> list[dict]:
    """Normalize RetrievedResult objects into SourceItem dicts."""
    out = []
    for r in results:
        out.append({
            "doc_id": r.doc_id,
            "ministry": r.metadata.get("ministry") or "-",
            "subject": r.metadata.get("subject") or "-",
            "date": r.metadata.get("date"),
            "document_type": r.metadata.get("document_type"),
            "score": float(r.score),
            "question": r.question,
            "answer": r.answer,
            "dense_score": float(r.dense_score) if r.dense_score is not None else None,
            "bm25_score": float(r.bm25_score) if r.bm25_score is not None else None,
            "rrf_score": float(r.rrf_score) if r.rrf_score is not None else None,
            "rerank_score": float(r.rerank_score) if r.rerank_score is not None else None,
        })
    return out


def _sse(obj: dict) -> str:
    return f"data: {json.dumps(obj)}\n\n"


# ─────────────────────────────────────────────────────────────────────────────
# Grounding verification (server-side)
# ─────────────────────────────────────────────────────────────────────────────
# After the answer streams, we extract every claim (proper nouns, acronyms,
# figures) and check it VERBATIM against the retrieved sources. Claims not
# found are reported — the frontend shows them as "not found in sources" so
# the user can see exactly what the model may have invented, and the grounding
# % badge reflects server-verified facts, not a browser heuristic.

import re as _re

_ACRONYM_STOPWORDS = {
    "THE", "AND", "FOR", "NOT", "ARE", "WAS", "WERE", "BUT", "HAS", "HAVE",
    "HAD", "ITS", "YOU", "OUR", "OUT", "OFF", "CAN", "MAY", "THIS", "THAT",
    "WITH", "FROM", "INTO", "WHEN", "WHAT", "WHY", "HOW", "THAN", "THEN",
    "INDIA", "GOVERNMENT", "MINISTRY", "ANSWER", "QUESTION", "STATE",
}

_FIGURE_RE = _re.compile(
    r"\b\d+(?:[.,]\d+)?\s?(?:%|mm|cm|km|m\b|MW|GW|KW|sq\.?\s?km|crore|lakh|"
    r"million|billion|hrs?|hours?|years?|deg(?:ree)?s?|₹|rs\.?)\b",
    _re.IGNORECASE,
)
# "48 Doppler Weather Radars", "32 Water Quality Buoys", "675 AWS" — number +
# a capitalized noun phrase (up to 4 words). Catches list-number swaps that a
# bare figure+unit regex misses ("32 Water Quality Buoys" vs the source's
# "2 Water Quality Buoys").
_NUM_WORD_RE = _re.compile(
    r"\b\d+(?:[.,]\d+)?\s+[A-Z][A-Za-z-]+(?:\s+[A-Z][A-Za-z-]+){0,3}\b"
)
_QUOTE_RE = _re.compile(r"\"([^\"\\]{6,80})\"")
_ACRONYM_RE = _re.compile(r"\b[A-Z]{2,8}\b")
_ACRONYM_PLURAL_RE = _re.compile(r"\b[A-Z]{2,7}[a-z]{1,2}\b")
_NAMED_ABBR_RE = _re.compile(r"\b[A-Z][A-Za-z&.\- ]{2,60}\s*\([A-Z]{2,10}\)")


def _extract_claims(answer: str, max_claims: int = 12) -> list[str]:
    """Extract a bounded set of checkable claims from the answer."""
    claims: list[str] = []
    for m in _FIGURE_RE.finditer(answer):
        claims.append(m.group(0).strip())
    for m in _NUM_WORD_RE.finditer(answer):
        claims.append(m.group(0).strip())
    for m in _QUOTE_RE.finditer(answer):
        claims.append(m.group(1).strip())
    for m in _NAMED_ABBR_RE.finditer(answer):
        claims.append(m.group(0).strip())
    for m in _ACRONYM_RE.finditer(answer):
        tok = m.group(0)
        if tok in _ACRONYM_STOPWORDS or len(tok) < 3:
            continue
        claims.append(tok)
    for m in _ACRONYM_PLURAL_RE.finditer(answer):
        tok = m.group(0)
        if tok in _ACRONYM_STOPWORDS or len(tok) < 3:
            continue
        claims.append(tok)
    # de-dup, keep order, cap
    seen: set[str] = set()
    out: list[str] = []
    for c in claims:
        key = c.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
        if len(out) >= max_claims:
            break
    return out


# ── Citation / grounding-aware filter (#3, non-destructive) ───────────────
# Original behavior dropped ANY sentence lacking "[Source N]", which cut
# correct-but-uncited prose and made answers look truncated. New behavior is
# claim-aware: a sentence is removed ONLY if it contains a checkable claim
# (figure, acronym, named phrase) that is NOT found in the retrieved sources —
# i.e. an actual hallucination risk. Plain prose, headings, and correctly
# grounded statements are never removed just for missing a citation token.

_CITATION_RE = _re.compile(r"\[\s*[Ss]ource\s*(\d+)\s*\]")


def _apply_citation_filter(
    answer: str,
    sources: list[dict],
    max_drop: int = 3,
) -> tuple[str, list[str]]:
    """Drop only sentences carrying UNVERIFIED claims. Returns
    (filtered_answer, dropped_sentences). Never empties the answer."""
    if not answer.strip() or not sources:
        return answer, []

    # Build a source haystack for per-sentence grounding.
    haystack = " ".join(
        f"{s.get('question','')} {s.get('answer','')}" for s in sources
    ).lower()

    import re as _re2
    sentences = _re2.split(r"(?<=[.!?])\s+|\n+", answer)
    kept: list[str] = []
    dropped: list[str] = []

    for s in sentences:
        s = s.strip()
        if not s:
            continue
        # A citation token always lets a sentence through.
        if _CITATION_RE.search(s):
            kept.append(s)
            continue
        # No claims to verify -> plain prose/connective -> keep.
        claims = _extract_claims(s)
        if not claims:
            kept.append(s)
            continue
        # Has claims but none are grounded -> hallucination risk -> drop.
        ungrounded = [c for c in claims if c.lower() not in haystack]
        if ungrounded and len(dropped) < max_drop:
            dropped.append(s)
            continue
        kept.append(s)

    if not kept:
        # never return an empty answer
        return answer, dropped
    return "\n\n".join(kept), dropped


# ── Normalized + alias-aware grounding (#1 recall fix) ────────────────────
# The raw verbatim check ("is the exact string in the source?") is precise but
# has poor recall: "Indian Space Research Organisation" vs "ISRO" is the same
# fact, but exact matching flags it. We keep the STRICT floor (a claim must
# have genuine textual support in a source) but make matching smarter:
#   1. normalize both sides (lowercase, unify currency/abbrevs, strip
#      punctuation, collapse whitespace),
#   2. resolve known alias groups (acronym <-> full name),
#   3. strip trailing plurals ("DWRs" -> "DWR").
# Invented names like "VSSC" stay NOT-FOUND because they are not in any alias
# group of a concept actually present in the sources.

_TOKEN_ALIASES = {
    "rs": "rupee", "inr": "rupee", "₹": "rupee",
    "&": "and", "ltd": "limited", "dept": "department",
    "govt": "government", "yr": "year", "hrs": "hours", "hr": "hour",
}

# Equivalent surface forms of the SAME entity/concept. These are true
# synonyms mined from the corpus — deliberately NOT including things like
# "VSSC" for ISRO (VSSC is a sub-entity; documents never use it for the
# personnel sphere, so it must keep failing the check).
# Equivalent surface forms of the same entity — SINGLE SOURCE of truth:
# frontend/src/utils/grounding_aliases.json. Editing the JSON is the ONLY way
# to change aliases; this list is loaded from it so backend and frontend can
# never drift again (was: two hardcoded copies, frontend missing 7 terms).
_ALIAS_GROUPS: list[list[str]] = json.loads(
    (PROJECT_ROOT / "frontend" / "src" / "utils"
     / "grounding_aliases.json").read_text(encoding="utf-8")
)["groups"]


def _normalize(text: str) -> str:
    """Lowercase, unify currency/abbrev tokens, strip punctuation, collapse."""
    t = text.lower()
    # remove digit-grouping commas first: "2,000" -> "2000" so it matches "2000"
    t = _re.sub(r"(?<=\d),(?=\d)", "", t)
    # token aliases with word boundaries, space-padded so they never glue to
    # neighbours ("₹2,000" -> "rupee 2000", "rs." -> "rupee")
    for k, v in _TOKEN_ALIASES.items():
        t = _re.sub(rf"\b{_re.escape(k)}\b", f" {v} ", t)
    # replace anything non-alphanumeric with a space
    t = _re.sub(r"[^a-z0-9]+", " ", t)
    return _re.sub(r"\s+", " ", t).strip()


def _singularize(norm: str) -> str:
    """Strip a trailing plural 's' if it leaves a meaningful token."""
    if len(norm) > 4 and norm.endswith("s") and not norm.endswith("ss"):
        return norm[:-1]
    return norm


def _claim_candidates(claim: str) -> list[str]:
    """All normalized surface forms that represent the same concept as the
    claim — the claim itself, its singular form, and every alias-group member
    if the claim names a known entity. Number+entity claims ("48 DWRs",
    "675 AWS") additionally expand the entity across its alias group and are
    matched with the number BEFORE or AFTER the entity (tables often list
    "AWS 675"), so a swapped figure ("32 Water Quality Buoys" vs the source's
    "2 Water Quality Buoys") still fails every candidate.
    """
    c = _normalize(claim)
    cands = [c, _singularize(c)]
    m = _re.match(r"^(\d+)\s+(.+)$", c)
    if m:
        num, phrase = m.group(1), m.group(2)
        phrases = {phrase, _singularize(phrase)}
        for group in _ALIAS_GROUPS:
            if any(mem in phrase or phrase in mem for mem in group if len(mem) > 2):
                phrases.update(group)
        for p in phrases:
            if p:
                cands.append(f"{num} {p}")
                cands.append(f"{p} {num}")
    else:
        for group in _ALIAS_GROUPS:
            if c in group or any(m in c for m in group if len(m) > 4):
                cands.extend(group)
    # de-dup
    seen: set[str] = set()
    out: list[str] = []
    for x in cands:
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def _claim_supported(claim: str, src_normalized: str) -> bool:
    """True if any surface form of the claim appears in the normalized source."""
    for cand in _claim_candidates(claim):
        if cand and cand in src_normalized:
            return True
    return False


def _grounding_report(answer: str, sources: list[dict]) -> list[dict]:
    """Verify each extracted claim against the sources using normalized +
    alias-aware matching. Returns [{text, found, source}]."""
    if not sources:
        return []
    claims = _extract_claims(answer)
    src_texts = [
        {
            "doc_id": s["doc_id"],
            "text": _normalize(f"{s.get('question','')} {s.get('answer','')}"),
        }
        for s in sources
    ]
    report: list[dict] = []
    for c in claims:
        found = False
        src = None
        for st in src_texts:
            if _claim_supported(c, st["text"]):
                found = True
                src = st["doc_id"]
                break
        report.append({"text": c, "found": found, "source": src})
    return report


# ── LLM judge (#2) ───────────────────────────────────────────────────────
# The regex grounding pass is fast but dumb: it only does verbatim substring
# matching, so a CITED-but-wrong claim (e.g. "VSSC" instead of "ISRO") slips
# through. The LLM judge re-verifies the claims the regex flagged as
# NOT-FOUND (and, optionally, the cited ones) by READING the sources and
# returning a supported/not-supported verdict per claim. It's a single short
# call over the flagged claims — NOT a regeneration of the whole answer — so
# latency stays near the generation cost (+10-20%, not 2x).

_JUDGE_JSON_RE = _re.compile(r"\{.*\}", _re.DOTALL)


def _llm_judge_claims(claims: list[dict], sources: list[dict]) -> list[dict]:
    """Verify flagged claims with an LLM judge against the sources.

    Returns the claims list with updated ``found`` / ``source`` / ``note``.
    On any failure (LLM offline, parse error) returns the original claims
    unchanged — the regex verdicts remain authoritative.
    """
    if not claims or not sources:
        return claims
    # Only judge claims the regex could NOT verify (the interesting ones).
    pending = [c for c in claims if not c.get("found")]
    if not pending:
        return claims

    try:
        # Compact source text: id + answer (truncated) per doc.
        src_blocks = []
        for i, s in enumerate(sources[:6], start=1):
            ans = (s.get("answer") or "")[:1500]
            src_blocks.append(f"[Source {i}] {s.get('doc_id','')}\n{ans}")
        src_text = "\n\n".join(src_blocks)

        claim_lines = "\n".join(
            f"{i}. {c['text']}" for i, c in enumerate(pending, start=1)
        )
        prompt = (
            "You are a strict evidence auditor. Below are source documents and a "
            "list of claims. For EACH claim decide whether it is SUPPORTED by the "
            "sources — supported means the claim's fact appears in the source text "
            "verbatim or is directly implied. If a claim names an organization, "
            "programme, or figure that does not appear in the sources, it is NOT "
            "supported.\n"
            "Return ONLY JSON (no prose):\n"
            '{"verdicts":[{"index":1,"supported":true,"source":"18-3-2571"},'
            '{"index":2,"supported":false,"source":null}]}\n\n'
            f"SOURCES:\n{src_text}\n\nCLAIMS:\n{claim_lines}"
        )
        resp = llm_client.generate(
            prompt=prompt,
            system=(
                "You are an evidence-verification assistant. Answer strictly with "
                "the requested JSON. Never invent claims or sources."
            ),
        )
        raw = resp.text
        m = _JUDGE_JSON_RE.search(raw or "")
        if not m:
            return claims
        data = json.loads(m.group(0))
        verdicts = data.get("verdicts") or []

        by_index = {}
        for v in verdicts:
            try:
                by_index[int(v.get("index"))] = v
            except (TypeError, ValueError):
                continue

        out = []
        pending_i = 0
        for c in claims:
            if c.get("found"):
                out.append(c)  # already verified verbatim by regex — keep
                continue
            pending_i += 1  # index within the pending list the judge saw
            v = by_index.get(pending_i)
            if v is None:
                out.append(c)  # judge gave no verdict — keep regex verdict
                continue
            supported = bool(v.get("supported"))
            src = v.get("source") or None
            # CRITICAL: the judge's "supported" is only trusted if it can name
            # a source that ACTUALLY contains the claim. The judge is the same
            # model that hallucinated the claim (e.g. "VSSC" — it "remembers"
            # VSSC is ISRO's space centre), so a bare "supported: true" must
            # NOT override a verbatim miss. Only accept the judge's verdict
            # when the claim text appears in the cited source's text (using the
            # same normalized+alias matcher as the grounding pass).
            judge_trusted = False
            if supported and src:
                for s in sources:
                    if s.get("doc_id") == src:
                        stxt = _normalize(
                            f"{s.get('question','')} {s.get('answer','')}"
                        )
                        if _claim_supported(c["text"], stxt):
                            judge_trusted = True
                        break
            if supported and not judge_trusted:
                # Judge said supported but cannot back it with verbatim text
                # in a real source → keep the regex verdict (rejected).
                out.append({
                    "text": c["text"],
                    "found": False,
                    "source": None,
                    "note": "rejected by LLM judge (no verbatim source support)",
                })
            else:
                out.append({
                    "text": c["text"],
                    "found": judge_trusted,
                    "source": src if judge_trusted else None,
                    "note": (
                        "verified by LLM judge" if judge_trusted
                        else "rejected by LLM judge"
                    ),
                })
        return out
    except Exception as e:  # noqa: BLE001 - judge must never break the stream
        import traceback
        print(f"[llm-judge] failed ({type(e).__name__}: {e}) — using regex verdicts")
        print(traceback.format_exc(limit=3))
        return claims


def _llm_rewrite_answer(
    answer: str,
    rejected_claims: list[str],
    sources: list[dict],
) -> str:
    """Second LLM call: rewrite the answer WITHOUT the judge-rejected claims.

    The rewrite keeps the structure and all supported content, but removes
    (or corrects) the unsupported statements. Returns the rewritten markdown
    answer. Raises on failure — the caller falls back to the original.
    """
    src_blocks = []
    for i, s in enumerate(sources[:6], start=1):
        ans = (s.get("answer") or "")[:1500]
        src_blocks.append(f"[Source {i}] {s.get('doc_id','')}\n{ans}")
    src_text = "\n\n".join(src_blocks)

    rejected_lines = "\n".join(f"- {c}" for c in rejected_claims)
    prompt = (
        "You are an evidence auditor. Below is a draft answer and a list of "
        "claims that were REJECTED because they are NOT supported by the source "
        "documents.\n\n"
        f"REJECTED CLAIMS:\n{rejected_lines}\n\n"
        f"DRAFT ANSWER:\n{answer}\n\n"
        f"SOURCES:\n{src_text}\n\n"
        "Rewrite the draft answer so that it:\n"
        "1. Removes every statement based on a rejected claim.\n"
        "2. Keeps all supported statements verbatim where possible.\n"
        "3. Does NOT add any new facts, names, or figures.\n"
        "4. Preserves markdown formatting and [Source N] citations for the "
        "kept statements.\n"
        "If everything was rejected, say the context does not support the "
        "claim.\n"
        "Return ONLY the rewritten answer, no commentary."
    )
    resp = llm_client.generate(
        prompt=prompt,
        system=(
            "You are an evidence-verification assistant. Rewrite the answer to "
            "remove unsupported claims. Never invent facts."
        ),
    )
    return (resp.text or "").strip()



def _remove_rejected_sentences(answer: str, rejected_claims: list[str]) -> tuple[str, list[str]]:
    """Remove sentences containing judge-rejected claims from the answer.

    Returns (cleaned_answer, removed_sentences). Only sentences that carry a
    claim the LLM judge explicitly rejected are removed — never plain prose,
    never grounded claims. This is the targeted enforcement that makes the
    visible answer correct, not just flagged.
    """
    if not answer.strip() or not rejected_claims:
        return answer, []
    import re as _re2
    sentences = _re2.split(r"(?<=[.!?])\s+|\n+", answer)
    kept: list[str] = []
    removed: list[str] = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        sl = s.lower()
        if any(rc.lower() in sl for rc in rejected_claims):
            removed.append(s)
        else:
            kept.append(s)
    if not kept:
        return answer, removed  # never empty the answer
    return "\n\n".join(kept), removed


def _resolve_exec(request: ChatStreamRequest):
    """Bind ACTIVE_CONFIG / llm_client / generator per execution profile —
    identical resolution path as chat_endpoint and the provider switch."""
    exec_mode = request.mode.lower()
    ACTIVE_CONFIG["mode"] = exec_mode
    _refresh_active_served_family()  # TTL-cached; no-op unless server provider
    family_id = ACTIVE_CONFIG["model_family"]
    family = model_registry.get(family_id) or model_registry.get("qwen2.5")
    plan, resolved_model = _apply_execution_plan(family, ACTIVE_CONFIG["provider"], exec_mode)
    llm_client.api_key = _active_api_key()
    print(f"[exec] mode={exec_mode} think={'ON' if llm_client.think else 'OFF'} "
          f"model={resolved_model} verify={ACTIVE_CONFIG['verify_depth']}")
    generator.context_budget_ratio = 0.80
    return family, plan, resolved_model


@app.post("/api/chat/stream")
def chat_stream(request: ChatStreamRequest):
    """SSE streaming endpoint for the drafting workspace.

    Emits: status (pipeline stages) -> sources -> trace -> tokens -> meta -> done.
    Never 500s: generation failures are surfaced as an ``error`` event with the
    retrieved sources still delivered.
    """
    def event_stream():
        query = request.message.strip()
        if not query:
            yield _sse({"type": "error", "message": "Query message cannot be empty"})
            yield _sse({"type": "done"})
            return

        ret_mode = request.retrieval_mode.lower()

        # ── User-Knowledge shortcut: same/similar question saved before? ──
        # If found, return the SAVED curated answer directly (no RAG, no LLM
        # call) — the user-curated answer always wins.
        try:
            hit = knowledge_lookup(query)
            if hit.get("found"):
                yield _sse({"type": "sources", "sources": hit.get("sources") or [], "is_graph": False})
                yield _sse({"type": "tokens", "text": hit.get("answer", "")})
                yield _sse({"type": "meta", "meta": {
                    "provider": "user-knowledge", "model": "saved-answer",
                    "profile": request.mode.lower(),
                    "retrieved_documents": len(hit.get("sources") or []),
                    "response_time_ms": 0.0, "is_fallback": False,
                    "knowledge_match": hit.get("matched"),
                }})
                yield _sse({"type": "status", "stage": "generate", "message": "Answered from saved knowledge", "done": True})
                yield _sse({"type": "phase", "phase": "done"})
                yield _sse({"type": "done"})
                return
        except Exception:  # noqa: BLE001 — lookup must never block a real query
            pass

        try:
            family, plan, resolved_model = _resolve_exec(request)
        except Exception as e:  # noqa: BLE001
            yield _sse({"type": "error", "message": f"Config resolution failed: {e}"})
            yield _sse({"type": "done"})
            return

        style_hint = ""
        original_system_prompt = generator.system_prompt
        # Tone is authoritative for register + verbosity. Frontend omits
        # draft_style for "Default Tone"; treat that as "default". Grounding
        # / citation rules stay in SYSTEM_PROMPT and are never replaced.
        style_key = (request.draft_style or "default").strip().lower()
        tone_templates = {
            "default": (
                "\n\nTONE: DEFAULT — write a complete, factual answer. "
                "Match length to what the question and sources require. "
                "Do not pad, and do not compress to a character quota. "
                "Prefer third-person official wording unless the sources "
                "read more naturally otherwise. Keep every [Source N] citation."
            ),
            "professional": (
                "\n\nTONE: PROFESSIONAL — write in a clear, formal, "
                "business-register style. Be precise, objective and "
                "confident. Use short structured paragraphs or labelled "
                "points. Avoid slang, hedging and first-person asides. "
                "Lead with the answer, then supporting detail. "
                "Length follows this register, not a character quota."
            ),
            "parliamentary": (
                "\n\nTONE: PARLIAMENTARY — mirror the style of a Lok Sabha "
                "ministry reply. Write in the third person: \"The Government "
                "has...\", \"As per available information...\", \"It may be "
                "stated that...\". Mirror the question's clauses as "
                "(a), (b), (c) sub-answers. Preserve figures, names and "
                "[Source N] citations verbatim. Stay factual and official; "
                "no opinions, no recommendations. Length may match a full "
                "ministry reply; do not cap at a character quota."
            ),
            "concise": (
                "\n\nTONE: CONCISE — be brief and direct. Use short bullet "
                "points or 1-2 sentence paragraphs. Give only the key facts "
                "and figures; drop elaboration, context and repetition. "
                "Keep every [Source N] citation."
            ),
            "detailed": (
                "\n\nTONE: DETAILED — give a comprehensive answer. Cover "
                "every aspect of the question with sub-sections or numbered "
                "points, include supporting context, figures, dates and "
                "institutional roles, and cite all relevant [Source N] "
                "references. Depth over brevity. Ignore any preference for "
                "short answers; stay grounded in the sources."
            ),
        }
        style_hint = tone_templates.get(
            style_key,
            f"\n\nTONE: Compose the answer in a {style_key} register. "
            "This TONE controls length and wording style.",
        )
        generator.system_prompt = (generator.system_prompt or "").rstrip() + style_hint

        t_ret_start = time.perf_counter()
        sources: list[dict] = []
        timings = None
        is_graph = (ret_mode == "graph" and not is_semantic_synthesis_query(query))

        try:
            if is_graph:
                # ── GRAPH / metadata traversal path ──
                stages = ["entities", "traversal", "expansion", "evidence"]
                for s in stages:
                    yield _sse({"type": "status", "stage": s, "message": s, "done": False})
                results = graph_retriever.retrieve(query, top_k=request.top_k)
                for s in stages:
                    yield _sse({"type": "status", "stage": s, "message": s, "done": True})
                sources = _to_sources(results)
                ret_latency = (time.perf_counter() - t_ret_start) * 1000
            else:
                # ── HYBRID RAG path with live stage callbacks ──
                stage_events: list[tuple[str, dict]] = []

                def _collect(name: str, info: dict) -> None:
                    stage_events.append((name, info))

                yield _sse({"type": "status", "stage": "embed", "message": "Embedding query…", "done": False})
                results, timings = pipeline.retrieve(
                    query, top_k=_effective_top_k(request.top_k, plan),
                    on_stage=_collect,
                    doc_types=request.doc_types,
                    orgs=request.orgs,
                    doc_categories=request.doc_categories,
                )
                # Task 3 (Deep): re-bond heading-like neighbor chunks
                _maybe_enrich_deep_neighbors(plan, results)
                for name, info in stage_events:
                    label = {
                        "embed": "Embed query", "dense": "Semantic search (dense)",
                        "bm25": "BM25 search", "rrf": "RRF fusion",
                        "rerank": "Reranking (cross-encoder)",
                    }.get(name, name)
                    yield _sse({
                        "type": "status", "stage": name, "message": label,
                        "count": info.get("count"), "done": True,
                    })
                sources = _to_sources(results)
                ret_latency = (time.perf_counter() - t_ret_start) * 1000

            # Task 3: emit EXACTLY what the model receives — the budget-admitted
            # set from the shared prepare_context cache (generate_stream reuses
            # the identical assembly). Not a fixed count: as many relevant
            # documents as the evidence budget fits.
            if not is_graph:
                admitted, adiag = _admission_diag(query, results)
                # Runtime evidence trace (Bug-2 diagnostics). Reads:
                #   CASE A  retrieved > admitted (budget) → correct Task-3
                #           behavior; skipped/evidence_used tell you why.
                #   CASE B  retrieved < retrieval_top_k → the retrieval side
                #           narrowed the pool (check dense/bm25/rrf/reranked
                #           counts); effective_top_k proves whether the plan
                #           pool (5 Standard / 10 Deep) reached the pipeline.
                #   CASE C  legacy_fallback=True → a planned request hit the
                #           max_context_docs cap — an alarm, never expected.
                # (stage_events exists here: this branch only runs on the
                # hybrid path, where it is populated.)
                stage_counts = {
                    name: info.get("count")
                    for name, info in stage_events
                    if "count" in info
                }
                print(
                    "[evidence-trace]"
                    f" mode={getattr(plan, 'mode', '?')}"
                    f" provider={ACTIVE_CONFIG['provider']}"
                    f" model={resolved_model}"
                    f" plan={'ON' if adiag.get('plan_attached') else 'MISSING'}"
                    f" retrieval_top_k={adiag.get('retrieval_top_k')}"
                    f" effective_top_k={_effective_top_k(request.top_k, plan)}"
                    f" dense={stage_counts.get('dense', '-')}"
                    f" bm25={stage_counts.get('bm25', '-')}"
                    f" rrf={stage_counts.get('rrf', '-')}"
                    f" reranked={stage_counts.get('rerank', '-')}"
                    f" retrieved={adiag.get('retrieved_count', adiag.get('pool', '-'))}"
                    f" admitted={adiag.get('admitted', len(admitted))}"
                    f" skipped={len(adiag.get('skipped_doc_ids') or [])}{adiag.get('skipped_doc_ids') or []}"
                    f" evidence_budget={adiag.get('plan_evidence_budget_tokens', '-')}"
                    f" evidence_budget_after_overhead={adiag.get('evidence_budget_tokens', '-')}"
                    f" evidence_used={adiag.get('evidence_used_tokens', '-')}"
                    f" legacy_max_context_docs_fallback={adiag.get('legacy_max_context_docs_fallback', '?')}"
                )
                sources = _filter_to_admitted(sources, admitted, key=lambda s: s["doc_id"])

            yield _sse({"type": "sources", "sources": sources, "is_graph": is_graph})

            trace_payload = None
            if not is_graph and timings is not None:
                trace_payload = {
                    "embed_query_ms": round(timings.embed_query_ms, 2),
                    "dense_search_ms": round(timings.dense_search_ms, 2),
                    "bm25_search_ms": round(timings.bm25_search_ms, 2),
                    "rrf_fusion_ms": round(timings.rrf_fusion_ms, 2),
                    "rerank_ms": round(timings.rerank_ms, 2),
                    "retrieval_total_ms": round(ret_latency, 2),
                    "truncated_docs": sorted(getattr(pipeline.reranker, "last_truncated_docs", set()) or set()),
                }
            if trace_payload:
                yield _sse({"type": "trace", "trace": trace_payload})

            if not sources:
                yield _sse({"type": "tokens", "text": (
                    "No relevant documents were retrieved from the knowledge base. "
                    "I cannot answer this question based on the available context."
                )})
                yield _sse({"type": "done"})
                return

            # ── Generation (streamed) ──
            yield _sse({"type": "status", "stage": "generate", "message": "Generating answer…", "done": False})
            # Tell the UI the model is now thinking (qwen3 emits reasoning
            # tokens live; qwen2.5 doesn't, so the UI shows a spinner + timer).
            yield _sse({"type": "phase", "phase": "thinking", "model": resolved_model})
            llm_available = llm_client.check_health(api_key=_active_api_key())
            if not llm_available:
                yield _sse({"type": "tokens", "text": (
                    "**[System Notice: LLM Generation Offline]**\n\n"
                    "The retrieved documents below were found, but the active LLM "
                    "service is offline or unauthorized. Start the local service or "
                    "switch provider, then retry."
                )})
                yield _sse({"type": "meta", "meta": {
                    "provider": ACTIVE_CONFIG["provider"], "model": resolved_model,
                    "profile": request.mode.lower(),
                    "retrieved_documents": len(sources), "retrieved_chunks": len(sources),
                    "response_time_ms": round((time.perf_counter() - t_ret_start) * 1000, 1),
                    "is_fallback": True,
                }})
                yield _sse({"type": "done"})
                return

            from src.retrieval.result import RetrievedResult
            context = [
                RetrievedResult(
                    doc_id=s["doc_id"], question=s["question"], answer=s["answer"],
                    score=s["score"], retrieval_method="hybrid", metadata={
                        "ministry": s["ministry"], "subject": s["subject"],
                        "document_type": s.get("document_type"),
                    },
                ) for s in sources
            ]
            # accumulate streamed text for the server-side grounding pass
            streamed_parts: list[str] = []
            for ev in generator.generate_stream(query, context):
                if ev["type"] == "tokens":
                    streamed_parts.append(ev["text"])
                    yield _sse({"type": "tokens", "text": ev["text"]})
                elif ev["type"] == "reasoning":
                    # Frontend boundary gate (Bug-1 invariant): reasoning may
                    # reach the canvas ONLY when the RESOLVED plan says
                    # thinking is ON. Standard resolves thinking=False — and
                    # some providers ignore think:false (several Ollama/qwen3
                    # builds) — so the application enforces the mode here,
                    # never the provider flag. The generator gates the same
                    # invariant upstream; this is the last boundary before
                    # the frontend either way.
                    if plan is None or plan.thinking:
                        # live chain-of-thought — forwarded to the Model Activity panel
                        yield _sse({"type": "reasoning", "text": ev["text"]})
                elif ev["type"] == "answer_start":
                    yield _sse({"type": "phase", "phase": "generating"})
                elif ev["type"] == "meta":
                    # generate_stream yields a FLAT meta event; wrap it for SSE
                    meta = {
                        "model": ev.get("model"),
                        "provider": ev.get("provider"),
                        "prompt_tokens": ev.get("prompt_tokens", 0),
                        "completion_tokens": ev.get("completion_tokens", 0),
                        "total_tokens": ev.get("total_tokens", 0),
                        "generation_latency_ms": ev.get("generation_latency_ms", 0.0),
                        "sources_used": ev.get("sources_used", []),
                    }
                    meta["provider"] = ACTIVE_CONFIG["provider"]
                    meta["profile"] = request.mode.lower()
                    meta["retrieved_documents"] = len(sources)
                    meta["retrieved_chunks"] = len(sources)
                    meta["response_time_ms"] = round(
                        (time.perf_counter() - t_ret_start) * 1000, 1
                    )
                    meta["is_fallback"] = False
                    yield _sse({"type": "meta", "meta": meta})

            # FAST PATH: emit the answer + done immediately. Verification
            # (grounding/judge/rewrite) runs AFTER the stream closes via
            # /api/verify — so the stream never blocks on extra LLM calls
            # (which caused the UI to appear "stuck" on slow local models).
            raw_text = "".join(streamed_parts)
            if not raw_text.strip():
                # Model produced reasoning (or nothing) but zero answer tokens —
                # surface WHY instead of a silent blank canvas. Usually the
                # generation hit the token limit during chain-of-thought.
                raw_text = (
                    "**[System Notice: model returned reasoning but no answer]**\n\n"
                    "The model produced its thinking but the final answer was empty. "
                    "This usually means the generation hit the token limit during "
                    "reasoning. Retry in Deep mode, or make the question more specific."
                )
            yield _sse({
                "type": "final",
                "text": raw_text,
                # Honest flag: these judge/citation values are streaming-time
                # placeholders (all zero). Real grounding/judge/rewrite runs
                # asynchronously via /api/verify after the stream closes.
                "verify_pending": True,
                "citation_dropped_count": 0,
                "citation_dropped": [],
                "judge_removed_count": 0,
                "judge_removed": [],
                "judge_rewritten": False,
            })
            yield _sse({"type": "status", "stage": "generate", "message": "Generating answer…", "done": True})
            yield _sse({"type": "phase", "phase": "done"})
            yield _sse({"type": "done"})

        except Exception as e:  # noqa: BLE001 - never let the stream 500
            import traceback
            print(f"[chat/stream] generation failed: {type(e).__name__}: {e}")
            print(traceback.format_exc(limit=5))
            if sources:
                yield _sse({"type": "sources", "sources": sources, "is_graph": is_graph})
            yield _sse({"type": "phase", "phase": "error"})
            yield _sse({"type": "error", "message": f"{type(e).__name__}: {str(e)[:300]}"})
            yield _sse({"type": "done"})
        finally:
            generator.system_prompt = original_system_prompt

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/api/verify")
def verify_answer(payload: dict):
    """Non-blocking post-generation verification.

    Runs AFTER the chat stream closes (so the stream never blocks on extra LLM
    calls). Returns the grounding report + a judge-rewritten answer with any
    unsupported claims removed. All LLM calls have timeouts so this can never
    hang.
    """
    answer = payload.get("answer") or ""
    sources = payload.get("sources") or []
    # Mode-aware depth: Fast = light (regex only, no LLM judge/rewrite —
    # instant). Deep = full (regex + LLM judge + rewrite).
    depth = payload.get("depth") or ACTIVE_CONFIG.get("verify_depth", "full")
    if not answer or not sources:
        return {"text": answer, "grounding": [], "judge_rewritten": False,
                "judge_removed_count": 0, "error": "missing answer or sources"}

    try:
        # 1. regex grounding
        grounding = _grounding_report(answer, sources)

        # 2. claim-aware filter (identify unverified sentences, informational)
        citation_dropped: list[str] = []
        if answer.strip():
            _filtered, citation_dropped = _apply_citation_filter(answer, sources)

        # 3. LLM judge — only in DEEP mode (Fast skips the extra LLM call)
        if depth == "full" and grounding and any(not c.get("found") for c in grounding):
            grounding = _llm_judge_claims(grounding, sources)

        # 4. rewrite — remove judge-rejected claims.
        # NOTE: the judge emits TWO rejection notes:
        #   "rejected by LLM judge"                                    (judge said unsupported)
        #   "rejected by LLM judge (no verbatim source support)"       (judge said supported but
        #                                                            couldn't back it verbatim —
        #                                                            the VSSC-type hallucination guard)
        # Both must be collected — use a prefix match so neither is missed.
        rejected_claims = [
            c["text"] for c in grounding
            if (not c.get("found"))
            and str(c.get("note", "")).startswith("rejected by LLM judge")
        ]
        final_text = answer
        judge_rewritten = False
        if depth == "full" and rejected_claims and answer.strip():
            rewrite = _llm_rewrite_answer(
                answer=answer, rejected_claims=rejected_claims, sources=sources,
            )
            if rewrite and rewrite.strip():
                final_text = rewrite
                judge_rewritten = True
                # deterministic safety net
                final_text, _ = _remove_rejected_sentences(final_text, rejected_claims)

        return {
            "text": final_text,
            "grounding": grounding,
            "judge_rewritten": judge_rewritten,
            "judge_removed_count": len(rejected_claims),
            "judge_removed": rejected_claims[:20],
            "citation_dropped_count": len(citation_dropped),
        }
    except Exception as e:  # noqa: BLE001 - never break the app
        import traceback
        print(f"[api/verify] failed ({type(e).__name__}: {e})")
        print(traceback.format_exc(limit=3))
        return {"text": answer, "grounding": [], "judge_rewritten": False,
                "judge_removed_count": 0, "error": f"{type(e).__name__}: {str(e)[:200]}"}


@app.post("/api/edit")
def ai_edit(payload: dict):
    """Stream an AI edit of the current draft (drafting workspace toolbar)."""
    instruction = (payload.get("instruction") or "").strip()
    document = payload.get("document") or ""
    style = (payload.get("draft_style") or "").strip()

    if not instruction:
        return StreamingResponse(iter([_sse({"type": "error", "message": "Instruction empty"}), _sse({"type": "done"})]),
                                 media_type="text/event-stream")
    if not document:
        return StreamingResponse(iter([_sse({"type": "error", "message": "No draft to edit"}), _sse({"type": "done"})]),
                                 media_type="text/event-stream")

    def event_stream():
        system = (
            "You are an AI editing assistant. Follow the user's instruction "
            "for the draft below — do whatever is asked (rewrite, restructure, "
            "summarize, expand, change tone, narrate, translate, etc.). "
            "Return the result in full, in the requested format.\n"
            "Grounding guardrail (only when the task is an EDIT of the draft's "
            "content, not a creative/formatting task): if you keep factual "
            "claims, preserve names, figures, and [Source N] citations as "
            "written; do not invent new facts. For creative/formatting tasks "
            "(story, style, structure), apply them freely.\n"
            f"DRAFT STYLE (if any): {style or 'none'}"
        )
        prompt = (
            f"EDIT INSTRUCTION:\n{instruction}\n\n"
            f"CURRENT DRAFT:\n{document}\n\n"
            f"Return the revised draft in full, markdown formatted."
        )
        prev_think = getattr(llm_client, "think", None)
        try:
            # Same provider/model/base_url as chat (ACTIVE_CONFIG / llm_client).
            # Thinking off: edits are rewrites, not Deep-mode reasoning.
            llm_client.think = False
            print(
                f"[edit] provider={llm_client.provider} model={llm_client.model} "
                f"base_url={getattr(llm_client, 'base_url', '')} think=OFF"
            )
            yield _sse({"type": "status", "stage": "edit",
                        "message": f"Editing with {llm_client.provider}/{llm_client.model}…",
                        "done": False})
            for chunk in llm_client.generate_stream(prompt=prompt, system=system):
                # Structured events (post client-stream migration): forward
                # only the visible tokens to the edit panel.
                if isinstance(chunk, dict):
                    if chunk.get("type") == "tokens":
                        yield _sse({"type": "tokens", "text": chunk.get("text", "")})
                    elif chunk.get("type") == "done":
                        break
                else:
                    yield _sse({"type": "tokens", "text": chunk})
            yield _sse({"type": "status", "stage": "edit", "message": "Editing with AI…", "done": True})
            yield _sse({"type": "done"})
        except Exception as e:  # noqa: BLE001
            yield _sse({"type": "error", "message": f"{type(e).__name__}: {str(e)[:300]}"})
            yield _sse({"type": "done"})
        finally:
            llm_client.think = prev_think

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─────────────────────────────────────────────────────────────────────────────
# Health endpoints (P1.8) — proxy/launcher liveness + readiness.
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/health/live")
def health_live():
    """Process is up (does not validate models/index)."""
    return {"status": "alive"}


@app.get("/health/ready")
def health_ready():
    """Readiness: index loaded + active LLM provider reachable.

    In APP_MODE=serve the serving process must be able to answer queries —
    if the index or LLM is missing, readiness fails so the orchestrator
    won't route traffic here (or won't call it 'up')."""
    ok_index = False
    try:
        p = pipeline._get()
        ok_index = len(p._doc_map) > 0 if not p.use_chunking else len(p._chunk_map) > 0
    except Exception:  # noqa: BLE001
        ok_index = False
    ok_llm = False
    try:
        ok_llm = llm_client.check_health(api_key=_active_api_key())
    except Exception:  # noqa: BLE001
        ok_llm = False
    ready = ok_index and ok_llm
    return {
        "status": "ready" if ready else "not_ready",
        "index_loaded": ok_index,
        "llm_healthy": ok_llm,
        "provider": ACTIVE_CONFIG.get("provider"),
        "model": ACTIVE_CONFIG.get("model"),
        "app_mode": os.environ.get("APP_MODE", "serve"),
    }


@app.get("/api/status")
def status():
    """Workstation header status: provider, model, mode, GPU.

    Additive discovery-sourced identity of the ACTIVE model (dynamic
    served-model discovery). Existing keys (provider/model_family/model/
    mode/retrieval_mode/gpu/enabled_providers) are unchanged — the frontend
    contract is preserved; all new keys are optional extras.
    """
    gpu = "CPU"
    try:
        import torch
        if torch.cuda.is_available():
            gpu = f"GPU ({torch.cuda.get_device_name(0)[:30]})"
    except Exception:  # noqa: BLE001
        gpu = "CPU"
    payload = {
        "provider": ACTIVE_CONFIG["provider"],
        "enabled_providers": _enabled_providers(),
        "model_family": ACTIVE_CONFIG["model_family"],
        "model": ACTIVE_CONFIG["model"],
        "mode": ACTIVE_CONFIG["mode"],
        "retrieval_mode": ACTIVE_CONFIG["retrieval_mode"],
        "gpu": gpu,
    }
    prov = ACTIVE_CONFIG["provider"]
    fam = model_registry.get(ACTIVE_CONFIG["model_family"])
    if fam is not None:
        payload["model_display_name"] = fam.display_name
        payload["model_metadata_source"] = fam.metadata_source  # catalog|server|fallback
        supported = fam.thinking.supported if fam.thinking else None
        payload["thinking_supported"] = supported  # true/false; null = UNKNOWN
    if prov in ("vllm", "openai_compatible"):
        prov_inst = provider_registry.get(prov)
        base = prov_inst._url() if prov_inst is not None else None  # env re-resolved
        if base:
            # in-cluster servers are unauthenticated; strip any userinfo
            # anyway so a credentialed env URL can never leak to the UI
            payload["provider_base_url"] = re.sub(r"(https?://)[^/@]+@", r"\1", base)
        payload["served_model"] = ACTIVE_CONFIG["model"]  # id sent on the wire
    if fam is not None:
        # Context identity, resolved by the SAME policy helper the budget
        # calculator uses — status/debug surface, additive keys only. None =
        # unknown (e.g. vLLM not exposing --max-model-len); nothing invented.
        from src.generation.policy import describe_context

        ctx = describe_context(
            catalog_context=int(fam.context_window),
            native_context=fam.native_context_window,
            serving_limit=_current_serving_limit(prov),
        )
        payload["native_context_tokens"] = ctx["native_context_tokens"]
        payload["serving_context_tokens"] = ctx["serving_context_tokens"]
        payload["app_context_limit_tokens"] = ctx["app_context_limit_tokens"]
        payload["effective_context_tokens"] = ctx["effective_context_tokens"]
    return payload


@app.post("/api/models/refresh")
async def refresh_models(provider: str | None = None):
    """Force served-model re-discovery (bypasses the TTL cache) — the manual
    refresh for HPC vLLM model swaps. The generation path picks the newly
    discovered active model up immediately (and would do so by itself at TTL
    expiry). A VLLM_MODEL pin keeps overriding whatever is discovered."""
    prov = (provider or ACTIVE_CONFIG["provider"]).lower().strip()
    if prov not in ("vllm", "openai_compatible"):
        return {"status": "skipped",
                "reason": f"provider {prov!r} is not served-model discovery-backed"}
    try:
        active = refresh_vllm_discovery()
    except VLLMDiscoveryError as e:
        raise HTTPException(status_code=503, detail=str(e))
    _refresh_active_served_family()  # apply to the running config now
    return {
        "status": "ok",
        "active_model": active["id"],
        "max_model_len": active.get("max_model_len"),
        "pinned": active["pinned"],
        "alternatives": active["alternatives"],
        "configured_model": ACTIVE_CONFIG["model"],
    }


# ─────────────────────────────────────────────────────────────────────────────
# Dynamic source catalogue — ministry tree (orgs) + doc categories + types
# present in the corpus, so the frontend source-filter renders whatever is
# actually present (new ministry/org appears automatically, no hardcoded list).
# ─────────────────────────────────────────────────────────────────────────────

_SOURCES_CACHE: dict = {"data": None, "key": None}


def _indexed_records():
    """Yield records from the *searchable* index, not the raw JSONL corpus."""
    inst = getattr(pipeline, "_instance", None)
    if inst is not None and getattr(inst, "_doc_map", None):
        return list(inst._doc_map.values())
    doc_map_path = resolve_index_dir() / "doc_map.json"
    if not doc_map_path.exists():
        return []
    import orjson

    data = orjson.loads(doc_map_path.read_bytes())
    return list(data.values()) if isinstance(data, dict) else []


@app.get("/api/sources")
def sources_catalogue():
    """Facets from the active searchable index (doc_map), not the raw corpus."""
    from src.retrieval.frontend.org_tree import build_sources_catalogue

    records = _indexed_records()
    inst = getattr(pipeline, "_instance", None)
    doc_map_path = resolve_index_dir() / "doc_map.json"
    key = (
        id(inst) if inst is not None else 0,
        len(records),
        doc_map_path.stat().st_mtime if doc_map_path.exists() else 0,
    )
    if _SOURCES_CACHE["key"] == key and _SOURCES_CACHE["data"] is not None:
        return _SOURCES_CACHE["data"]
    data = build_sources_catalogue(records)
    _SOURCES_CACHE.update({"data": data, "key": key})
    return data


# ─────────────────────────────────────────────────────────────────────────────
# User-Knowledge API — curated Q&A saved by scientists
# ─────────────────────────────────────────────────────────────────────────────
# user-knowledge/ holds {question, answer, sources, saved_at} JSON files.
# On a NEW query the system FIRST checks this folder for an exact/normalized
# question match; if found, the SAVED answer is returned directly (the
# user-curated answer always wins — no RAG, no hallucination). No embeddings
# needed for this lookup.

USER_KNOWLEDGE_DIR = user_knowledge_dir()


def knowledge_fuzzy_threshold() -> float:
    """Documented default 0.85. Override: KNOWLEDGE_FUZZY_THRESHOLD."""
    raw = (os.environ.get("KNOWLEDGE_FUZZY_THRESHOLD") or "0.85").strip()
    try:
        val = float(raw)
    except ValueError:
        return 0.85
    if val < 0.0 or val > 1.0:
        return 0.85
    return val


def _normalize_q(text: str) -> str:
    """Lowercase, strip punctuation/whitespace for question matching."""
    import re as _re

    t = text.lower()
    t = _re.sub(r"[^a-z0-9]+", " ", t)
    return _re.sub(r"\s+", " ", t).strip()


def _load_user_knowledge() -> list[dict]:
    """All saved knowledge entries (question normalized for matching)."""
    entries = []
    if not USER_KNOWLEDGE_DIR.exists():
        return entries
    for f in sorted(USER_KNOWLEDGE_DIR.glob("*.json")):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            data["_file"] = f.name
            data["_q_norm"] = _normalize_q(str(data.get("question", "")))
            entries.append(data)
        except Exception:  # noqa: BLE001
            continue
    return entries


@app.get("/api/knowledge-lookup")
def knowledge_lookup(q: str):
    """Find a saved answer for a question. Exact normalized match first,
    then fuzzy (difflib ratio >= 0.85) fallback. Returns {found, answer,
    sources, question} or {found: false}."""
    if not q or not q.strip():
        return {"found": False}
    qn = _normalize_q(q)
    entries = _load_user_knowledge()
    if not entries:
        return {"found": False}

    # exact normalized match
    for e in entries:
        if e["_q_norm"] == qn:
            return {"found": True, "answer": e.get("answer", ""),
                    "sources": e.get("sources", []), "question": e.get("question", q),
                    "matched": "exact", "saved_by": e.get("saved_by")}
    # fuzzy fallback (similar wording, e.g. "Doppler Radars" vs
    # "Doppler Weather Radars" ~0.91). Default 0.85 — override with
    # KNOWLEDGE_FUZZY_THRESHOLD (must stay in 0..1).
    import difflib

    threshold = knowledge_fuzzy_threshold()
    best, best_ratio = None, 0.0
    for e in entries:
        r = difflib.SequenceMatcher(None, qn, e["_q_norm"]).ratio()
        if r > best_ratio:
            best, best_ratio = e, r
    if best and best_ratio >= threshold:
        return {"found": True, "answer": best.get("answer", ""),
                "sources": best.get("sources", []), "question": best.get("question", q),
                "matched": "fuzzy", "score": round(best_ratio, 3),
                "saved_by": best.get("saved_by")}
    return {"found": False}


def _resolve_saved_by(request: Request | None, payload: dict | None = None) -> str:
    """Local identity for saved_by (no SSO). Priority: payload, X-User, APP_USER, OS user."""
    if payload:
        raw = (payload.get("saved_by") or "").strip()
        if raw:
            return raw[:128]
    if request is not None:
        hdr = (request.headers.get("X-User") or request.headers.get("X-Forwarded-User") or "").strip()
        if hdr:
            return hdr[:128]
    env = (os.environ.get("APP_USER") or os.environ.get("USERNAME") or os.environ.get("USER") or "").strip()
    if env:
        return env[:128]
    try:
        import getpass

        return (getpass.getuser() or "local-user")[:128]
    except Exception:
        return "local-user"


@app.post("/api/save-knowledge")
def save_knowledge(payload: dict, request: Request):
    """Save a curated Q&A into user-knowledge/<slug>.json. Overwrites if the
    same question was saved before (so re-saving an edited answer updates it)."""
    question = (payload.get("question") or "").strip()
    answer = (payload.get("answer") or "").strip()
    if not question or not answer:
        raise HTTPException(status_code=400, detail="question and answer required")
    import re as _re

    USER_KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    slug = _re.sub(r"[^a-z0-9]+", "_", question.lower())[:60] or "knowledge"
    dest = USER_KNOWLEDGE_DIR / f"{slug}.json"
    # Trim sources to citation identity ONLY (doc_id/subject/ministry/type) —
    # the full document texts are already in the main index; storing them here
    # would bloat the file and slow every lookup.
    def _trim(s: dict) -> dict:
        return {
            "doc_id": s.get("doc_id") or "",
            "subject": s.get("subject") or "",
            "ministry": s.get("ministry") or "",
            "document_type": s.get("document_type") or "",
            "score": s.get("score"),
        }
    trimmed_sources = [_trim(s) for s in (payload.get("sources") or []) if isinstance(s, dict)]
    # update existing file (don't create duplicates for same question)
    saved_by = _resolve_saved_by(request, payload)
    entry = {
        "question": question,
        "answer": answer,
        "sources": trimmed_sources,
        "saved_at": datetime.now().isoformat(timespec="seconds"),
        "saved_by": saved_by,
    }
    dest.write_text(json.dumps(entry, ensure_ascii=False, indent=1), encoding="utf-8")
    print(f"[save-knowledge] {dest.name} by {saved_by} ({len(trimmed_sources)} sources trimmed)")
    return {"status": "saved", "file": dest.name, "saved_by": saved_by}


# ─────────────────────────────────────────────────────────────────────────────
# Ingest API — scientists / sir can add internal documents via the UI
# ─────────────────────────────────────────────────────────────────────────────
_INGEST_STATE: dict = {
    "running": False,
    "last": None,          # {"at": iso, "ok": n, "failed": n, "records": n, "message": str}
    "pending": 0,
}


@app.get("/api/ingest/status")
def ingest_status():
    """How many files are waiting in the inbox + whether ingest is running."""
    inbox = inbox_dir()
    pending = 0
    if inbox.exists():
        pending = sum(1 for p in inbox.iterdir() if p.is_file())
    _INGEST_STATE["pending"] = pending
    return {
        "running": _INGEST_STATE["running"],
        "pending": pending,
        "last": _INGEST_STATE["last"],
        "inbox": str(inbox),
    }


def _run_inbox_ingest() -> None:
    """Append inbox files to the corpus, then rebuild + swap the live index."""
    _INGEST_STATE["running"] = True
    try:
        # 1. Convert inbox files IN-PROCESS (same tested ingest_folder path).
        #    No subprocess: conversion errors show up directly, and there is
        #    no cwd/env ambiguity between where uploads land and where the
        #    converter looks.
        from src.scripts.ingest_folder import ingest_folder as _ingest_folder
        import src.scripts.ingest_folder as _ingest_folder_mod

        # pin the converter to APP_* / project-root paths (never CWD)
        _ingest_folder_mod.CORPUS = corpus_path()
        _ingest_folder_mod.LOG = data_dir() / "sync.log"
        _ingest_folder_mod.INDEX_DIR = str(resolve_index_dir())
        _ingest_folder_mod.CORPUS.parent.mkdir(parents=True, exist_ok=True)

        inbox = inbox_dir()
        inbox.mkdir(parents=True, exist_ok=True)
        conv = _ingest_folder(str(inbox), move_processed=True)
        print(f"[ingest] conversion: {conv}")
        ok = conv.get("files", 0)
        added_count = conv.get("added", 0)
        failed = conv.get("failed", 0)
        # 2. INCREMENTALLY update the index (embed only NEW records) and swap
        #    the live pipeline — no full re-embed of the whole corpus.
        try:
            from src.data.loader import DataLoader
            from src.retrieval.hybrid.pipeline import HybridRAGPipeline
            import os as _os2

            corpus = corpus_path()
            idx_dir = resolve_index_dir()
            _ingest_embedded = 0  # how many new vectors went into the index
            if corpus.exists():
                records = DataLoader.load_jsonl(corpus)
                _env2 = dict(_os2.environ)
                _env2["PYTHONIOENCODING"] = "utf-8"
                if all((idx_dir / f).exists() for f in (
                        "vector_store.index", "bm25_index.pkl", "doc_map.json", "pipeline_metadata.json")):
                    # existing index -> incremental (fast)
                    new_pipe = HybridRAGPipeline()
                    new_pipe.load(str(idx_dir))
                    n = new_pipe.add_records(records)
                    _ingest_embedded = int(n or 0)
                    if n:
                        new_pipe.save(str(idx_dir))
                    print(f"[ingest] incremental update: {n} new record(s) embedded+added")
                else:
                    # no index yet -> full build
                    new_pipe = HybridRAGPipeline(records=records)
                    new_pipe.save(str(idx_dir))
                    _ingest_embedded = len(records)
                    print(f"[ingest] full build with {len(records):,} records")
                pipeline.swap(new_pipe)
                _SOURCES_CACHE.update({"data": None, "key": None})
                print(f"[ingest] embeddings done: {_ingest_embedded} new vector(s) in live index")
        except Exception as e:  # noqa: BLE001
            print(f"[ingest] index rebuild failed: {e}")
            _INGEST_STATE["last"] = {
                "at": datetime.now().isoformat(timespec="seconds"),
                "ok": 0, "failed": 0, "records": 0,
                "message": f"Corpus updated but index rebuild failed: {e}. Restart server to reload.",
            }
            return

        _INGEST_STATE["last"] = {
            "at": datetime.now().isoformat(timespec="seconds"),
            "ok": ok, "failed": failed, "records": added_count,
            "message": f"{ok} file(s) ingested, {added_count} record(s) added, "
                       f"{_ingest_embedded} new record(s) embedded & indexed.",
        }
    except Exception as e:  # noqa: BLE001
        _INGEST_STATE["last"] = {
            "at": datetime.now().isoformat(timespec="seconds"),
            "ok": 0, "failed": 0, "records": 0,
            "message": f"Ingest failed: {e}",
        }
    finally:
        _INGEST_STATE["running"] = False


def _serve_mode_blocked():
    """In APP_MODE=serve the serving container must not mutate corpus/index."""
    if os.environ.get("APP_MODE", "").strip().lower() == "serve":
        raise HTTPException(
            status_code=403,
            detail="Read-only serve mode: ingestion is disabled. Run the ingest/build "
                   "job container instead (APP_MODE=ingest).",
        )
    return None


@app.post("/api/ingest")
def ingest_trigger(payload: dict):
    """Trigger ingestion of new files in data/inbox (background thread)."""
    _serve_mode_blocked()  # no-op unless APP_MODE=serve
    if _INGEST_STATE["running"]:
        return {"status": "busy"}
    _threading.Thread(target=_run_inbox_ingest, daemon=True).start()
    return {"status": "started"}


@app.post("/api/upload")
async def upload_document(request: Request):
    """Save an uploaded document (PDF/txt/json/jsonl) into data/inbox/.

    Raw body upload (application/octet-stream) — no python-multipart needed.
    Filename comes as a query param: POST /api/upload?filename=report.pdf
    Returns the saved file info; then call POST /api/ingest to process.
    """
    _serve_mode_blocked()  # no-op unless APP_MODE=serve
    filename = request.query_params.get("filename", "").strip()
    if not filename:
        raise HTTPException(status_code=400, detail="filename query param required")
    name = os.path.basename(filename)  # strip any path traversal
    if not name:
        raise HTTPException(status_code=400, detail="invalid filename")

    ext = Path(name).suffix.lower()
    if ext not in (".pdf", ".txt", ".md", ".json", ".jsonl"):
        raise HTTPException(status_code=400, detail=f"Unsupported file type {ext}")

    # Size guard BEFORE reading the body into RAM — a 2GB file would OOM the
    # worker (HPC shared nodes). Reject over-large uploads up front via the
    # Content-Length header, then cap the actual read as a second check.
    MAX_UPLOAD_BYTES = 200 * 1024 * 1024  # 200 MB
    cl = request.headers.get("content-length")
    if cl and cl.isdigit() and int(cl) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large (max {MAX_UPLOAD_BYTES // (1024 * 1024)} MB)",
        )

    body = await request.body()
    if len(body) < 10:
        raise HTTPException(status_code=400, detail="File is empty")
    if len(body) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large (max {MAX_UPLOAD_BYTES // (1024 * 1024)} MB)",
        )

    inbox = inbox_dir()
    inbox.mkdir(parents=True, exist_ok=True)
    dest = inbox / name
    dest.write_bytes(body)
    print(f"[upload] saved {name} ({len(body)} bytes) -> {dest}")
    return {"status": "saved", "file": name, "size": len(body)}


@app.get("/api/graph/build-status")
def graph_build_status():
    """Live Graph build progress read from the checkpoint file (if any)."""
    cp = resolve_graph_dir() / "checkpoint.json"
    if not cp.exists():
        return {"running": False, "documents_processed": 0, "failed": 0,
                "last_updated": None, "checkpoint_exists": False}
    try:
        data = json.loads(cp.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return {"running": False, "documents_processed": 0, "failed": 0,
                "last_updated": None, "checkpoint_exists": True}
    docs = data.get("documents", {}) if isinstance(data, dict) else {}
    done = sum(1 for v in docs.values() if v.get("status") == "done")
    failed = sum(1 for v in docs.values() if v.get("status") == "failed")
    return {
        "running": False,
        "documents_processed": done,
        "failed": failed,
        "total": len(docs),
        "last_updated": cp.stat().st_mtime,
        "checkpoint_exists": True,
        "path": str(cp),
    }


@app.post("/api/export")
def export_document(payload: dict):
    """Export the current answer as markdown / txt / docx."""
    fmt = (payload.get("format") or "md").lower()
    title = payload.get("title") or "answer"
    content = payload.get("content") or ""

    if fmt == "docx":
        try:
            from docx import Document  # python-docx
            doc = Document()
            doc.add_heading(title, level=1)
            for para in content.split("\n"):
                line = para.strip()
                if not line:
                    continue
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return Response(
                content=buf.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
            )
        except ImportError:
            raise HTTPException(
                status_code=409,
                detail="python-docx is not installed on this machine. Use format=md instead.",
            )
    if fmt == "txt":
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{title}.txt"'},
        )
    # default: markdown
    return Response(
        content=content,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{title}.md"'},
    )


@app.get("/{path:path}", response_class=HTMLResponse)
async def spa_history_fallback(path: str):
    """SPA history fallback: hard refresh on /workspace, /settings, etc.

    Must not swallow /api, /health, or /assets (those have their own routes;
    unknown paths under those prefixes stay 404).
    """
    first = (path.split("/", 1)[0] or "").lower()
    if first in _SPA_RESERVED_PREFIXES:
        raise HTTPException(status_code=404, detail="Not Found")
    return _serve_spa_index()


def start_server(port: int = 8000) -> None:
    """Run the FastAPI application on host 0.0.0.0.

    Smoke-test / first container deployment is **single-worker only**.
    ACTIVE_CONFIG, llm_client, and generator are in-process globals — do not
    pass workers>1 (or --workers) until that architecture is redesigned.
    """
    uvicorn.run(app, host="0.0.0.0", port=port, workers=1)


if __name__ == "__main__":
    start_server()
